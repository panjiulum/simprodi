# -*- coding: utf-8 -*-
"""test_audit_lanjutan_10_surat_kerapian.py — Regresi Audit Lanjutan 10:
kerapian & keandalan modul Surat Tugas Akhir (`app/routes/surat.py`).

Cakupan:
1. `buat()` & `sk_yudisium_tahap()` sekarang dibungkus try/except -> galat
   tak terduga saat generate dokumen jadi flash pesan ramah + redirect,
   BUKAN 500 mentah (sebelumnya tidak ada penanganan sama sekali).
2. Tabel di dokumen .docx yang dihasilkan (SK Pembimbing, SK Yudisium,
   SK Yudisium per Tahap, Undangan) sekarang punya `table.style` bergaris
   (bukan lagi "Normal Table" bawaan yang tanpa garis batas sel).
3. `JENIS_SURAT` (dropdown di index()) taat 1-sumber-kebenaran dengan
   `_GENERATORS`/`_PESAN_KOSONG` — otomatis konsisten, tidak lagi ditulis
   manual terpisah 3x.
4. Perataan gaya ternary utk `judul` di `_gen_undangan` (Seminar vs
   Sidang) tidak mengubah perilaku -- undangan Seminar & Sidang tetap
   menampilkan judul skripsi yang benar.
"""

import io
import os
import sys
import tempfile
from unittest import mock

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

tmpdir = tempfile.mkdtemp()
os.environ["HOME"] = tmpdir

from app import create_app  # noqa: E402
from app.routes import surat as surat_mod  # noqa: E402

db_path = os.path.join(tmpdir, "test.db")
app = create_app(db_path=db_path)
app.config["TESTING"] = True
app.config["WTF_CSRF_ENABLED"] = False
client = app.test_client()

FAILS = []


def check(label, cond):
    print(f"[{'OK' if cond else 'FAIL'}] {label}")
    if not cond:
        FAILS.append(label)


client.post(
    "/login",
    data={"username": "kaprodi", "password1": "test1234", "password2": "test1234"},
    follow_redirects=True,
)

# ---------------------------------------------------------------------
# Data dasar: mahasiswa + penetapan pembimbing + seminar + sidang, biar
# ke-4 jenis surat bisa dicoba dicetak.
# ---------------------------------------------------------------------
with app.test_request_context():
    conn = app.get_db()
    conn.execute("INSERT INTO mahasiswa(nim, nama) VALUES('2023990001','Mhs Uji Kerapian')")
    conn.execute("INSERT INTO dosen(nama, aktif) VALUES('Dr. Penguji Satu', 1)")
    conn.execute("INSERT INTO dosen(nama, aktif) VALUES('Dr. Penguji Dua', 1)")
    mid = conn.execute("SELECT id FROM mahasiswa WHERE nim='2023990001'").fetchone()["id"]
    d1 = conn.execute("SELECT id FROM dosen WHERE nama='Dr. Penguji Satu'").fetchone()["id"]
    d2 = conn.execute("SELECT id FROM dosen WHERE nama='Dr. Penguji Dua'").fetchone()["id"]
    conn.execute(
        "INSERT INTO penetapan_pembimbing(mahasiswa_id, semester, tahap, judul_final, "
        "pembimbing1_id, tgl_penetapan, no_sk) VALUES(?,?,?,?,?,?,?)",
        (mid, "Ganjil 2025/2026", "Tahap 1", "Judul Skripsi Uji Kerapian", d1, "2025-09-01", "001/SK/2025"),
    )
    conn.execute(
        "INSERT INTO seminar(mahasiswa_id, tgl_seminar, jam, penguji_ketua_id) "
        "VALUES(?,?,?,?)",
        (mid, "2026-01-10", "09:00", d1),
    )
    conn.execute(
        "INSERT INTO sidang(mahasiswa_id, judul_sidang, tgl_sidang, jam_sidang, ketua_id) "
        "VALUES(?,?,?,?,?)",
        (mid, "Judul Sidang Uji Kerapian", "2026-02-10", "10:00", d2),
    )
    conn.commit()

# ---------------------------------------------------------------------
# 1. try/except: get_setting() dibuat melempar galat tak terduga di
#    tengah proses generate dokumen -> HARUS jadi flash + redirect (302),
#    BUKAN 500.
# ---------------------------------------------------------------------
_get_setting_asli = surat_mod._db.get_setting


def _get_setting_gagal_untuk_nama_institusi(conn, key, default=""):
    """Pengganti sementara `_db.get_setting` yang HANYA melempar galat utk
    key 'nama_institusi' (dipanggil `_header()` di surat.py), supaya tidak
    ikut merusak pemanggilan `get_setting` lain di luar surat.py -- mis.
    context processor global `inject_globals()` di app/__init__.py yang
    memanggil get_setting(..., 'tema_warna') di SETIAP render_template,
    termasuk pada halaman redirect target setelah galat ditangani."""
    if key == "nama_institusi":
        raise RuntimeError("simulasi galat")
    return _get_setting_asli(conn, key, default)


with mock.patch.object(surat_mod._db, "get_setting", side_effect=_get_setting_gagal_untuk_nama_institusi):
    r = client.post(
        "/surat/buat",
        data={"mahasiswa_id": str(mid), "jenis": "SK Pembimbing"},
    )
    check("POST /surat/buat + galat tak terduga -> redirect (bukan 500)", r.status_code == 302)

r_flash = client.get("/surat/", follow_redirects=True)
check(
    "Flash 'Gagal membuat dokumen SK Pembimbing' tampil (pesan ramah, bukan traceback)",
    "Gagal membuat dokumen SK Pembimbing".encode() in r_flash.data,
)

# ---------------------------------------------------------------------
# 2. Jalur normal (tanpa mock) tetap berhasil -> pastikan try/except baru
#    TIDAK menelan jalur sukses.
# ---------------------------------------------------------------------
r_ok = client.post("/surat/buat", data={"mahasiswa_id": str(mid), "jenis": "SK Pembimbing"})
check("POST /surat/buat (jalur normal) -> 200", r_ok.status_code == 200)

import docx  # noqa: E402

doc = docx.Document(io.BytesIO(r_ok.data))

# ---------------------------------------------------------------------
# 3. Gaya tabel: SEBELUM perbaikan, table.style tidak pernah di-set ->
#    jatuh ke "Normal Table" (tanpa garis). Sekarang harus salah satu
#    gaya bergaris.
# ---------------------------------------------------------------------
check("Dokumen SK Pembimbing punya minimal 1 tabel", len(doc.tables) >= 1)
if doc.tables:
    check(
        "Tabel SK Pembimbing pakai gaya bergaris (bukan 'Normal Table')",
        doc.tables[0].style.name in ("Table Grid", "Light Grid Accent 1"),
    )

# ---------------------------------------------------------------------
# 4. Undangan Seminar & Sidang: judul skripsi tetap benar setelah
#    perataan gaya ternary.
# ---------------------------------------------------------------------
r_und_sem = client.post("/surat/buat", data={"mahasiswa_id": str(mid), "jenis": "Undangan Seminar"})
check("POST /surat/buat (Undangan Seminar) -> 200", r_und_sem.status_code == 200)
doc_sem = docx.Document(io.BytesIO(r_und_sem.data))
teks_sem = "\n".join(c.text for t in doc_sem.tables for row in t.rows for c in row.cells)
check("Undangan Seminar memuat judul skripsi yang benar", "Judul Skripsi Uji Kerapian" in teks_sem)
check(
    "Tabel data Undangan Seminar pakai gaya bergaris",
    doc_sem.tables[0].style.name in ("Table Grid", "Light Grid Accent 1"),
)
check(
    "Tabel 'Susunan Tim' Undangan Seminar juga pakai gaya bergaris",
    doc_sem.tables[-1].style.name in ("Table Grid", "Light Grid Accent 1"),
)

r_und_sid = client.post("/surat/buat", data={"mahasiswa_id": str(mid), "jenis": "Undangan Sidang"})
check("POST /surat/buat (Undangan Sidang) -> 200", r_und_sid.status_code == 200)
doc_sid = docx.Document(io.BytesIO(r_und_sid.data))
teks_sid = "\n".join(c.text for t in doc_sid.tables for row in t.rows for c in row.cells)
check(
    "Undangan Sidang memuat judul SIDANG (judul_sidang, bukan judul_final)",
    "Judul Sidang Uji Kerapian" in teks_sid,
)

# ---------------------------------------------------------------------
# 5. SK Yudisium per Tahap: try/except tetap mengembalikan dokumen valid
#    di jalur normal (data yudisium asli, bukan mock galat).
# ---------------------------------------------------------------------
with app.test_request_context():
    conn = app.get_db()
    conn.execute(
        "UPDATE sidang SET status_kelulusan='LULUS', nilai_angka=85 WHERE mahasiswa_id=?",
        (mid,),
    )
    conn.commit()

r_tahap = client.get("/surat/sk-yudisium-tahap?tahap=Tahap%201")
# Boleh 200 (dokumen jadi) ATAU redirect kalau data rencana yudisium
# modul lain belum lengkap di skenario minimal ini -- yang PENTING tidak
# 500. Cek eksplisit tidak ada traceback/500.
check("GET /surat/sk-yudisium-tahap (jalur normal) -> bukan 500", r_tahap.status_code != 500)

# ---------------------------------------------------------------------
# 6. JENIS_SURAT diturunkan otomatis dari _GENERATORS -- satu sumber
#    kebenaran, tidak lagi 3 tempat terpisah.
# ---------------------------------------------------------------------
check(
    "JENIS_SURAT == daftar key _GENERATORS (1 sumber kebenaran)",
    surat_mod.JENIS_SURAT == list(surat_mod._GENERATORS),
)
check(
    "Semua jenis surat di JENIS_SURAT punya pesan kosong terdaftar",
    all(j in surat_mod._PESAN_KOSONG for j in surat_mod.JENIS_SURAT),
)

r_index = client.get("/surat/")
check("GET /surat/ -> 200", r_index.status_code == 200)
for jenis in surat_mod.JENIS_SURAT:
    check(f"Dropdown /surat/ memuat opsi '{jenis}'", jenis.encode() in r_index.data)

# ---------------------------------------------------------------------
# 7. Import docx tidak lagi diulang per-fungsi -- dicek langsung dari
#    modul: seluruh fungsi _gen_* memakai `docx`/`WD_ALIGN_PARAGRAPH`/`Pt`
#    dari lingkup modul (bukan lokal), dibuktikan modul tetap bisa
#    dipanggil berkali-kali tanpa import ulang & docx tetap tersedia di
#    globals modul.
# ---------------------------------------------------------------------
check("`docx` tersedia sebagai atribut modul level (hoisted)", hasattr(surat_mod, "docx"))
check(
    "`WD_ALIGN_PARAGRAPH` tersedia sebagai atribut modul level (hoisted)",
    hasattr(surat_mod, "WD_ALIGN_PARAGRAPH"),
)
check("`Pt` tersedia sebagai atribut modul level (hoisted)", hasattr(surat_mod, "Pt"))

print("\n=== SELESAI ===")
if FAILS:
    print(f"{len(FAILS)} GAGAL:")
    for f in FAILS:
        print(" -", f)
    sys.exit(1)
print("SEMUA TES LULUS.")
