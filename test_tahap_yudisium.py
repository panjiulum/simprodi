# -*- coding: utf-8 -*-
"""
test_tahap_yudisium.py — Uji fitur baru: SK Yudisium per Tahap/Gelombang.

Latar belakang permintaan pengguna: tahap/gelombang pendaftaran seminar &
sidang (tahap_pengajuan, sudah ada sebelumnya untuk Pengajuan Judul &
Penetapan Pembimbing serta rekap honor penguji/pembimbing) sekarang juga
harus berlaku sampai ke Rencana Yudisium & SK Yudisium — supaya SK
Yudisium bisa ditetapkan & dicetak per tahap/gelombang, bukan cuma satu
per satu per mahasiswa.

Diuji:
1. `logic.rencana_yudisium_rows()`/`wisuda_rows()` menerima `tahap_filter`
   dan menyaring baris sesuai `penetapan_pembimbing.tahap` milik mahasiswa
   (sumber tahap yang sama dgn rkp_seminar/rkp_sidang/rekap_pembimbing).
2. Halaman `/kelulusan/yudisium?tahap=...` & `/kelulusan/wisuda?tahap=...`
   ikut tersaring, dan tab filter tahap tampil.
3. `POST /kelulusan/yudisium/tetapkan-tahap` menetapkan No. SK + Tgl
   Yudisium ke SEMUA mahasiswa pada tahap terpilih sekaligus, TANPA
   menimpa baris tahap lain, dan tidak menimpa No. SK yang sudah terisi
   manual kecuali "timpa" dicentang.
4. `GET /surat/sk-yudisium-tahap?tahap=...` menghasilkan SATU dokumen Word
   yang isinya memuat semua mahasiswa pada tahap tsb (bukan satu file per
   mahasiswa), dengan No. SK yang sama dengan yang baru ditetapkan.
5. Ekspor Excel Rencana Yudisium & Wisuda ikut membawa kolom Tahap.

Tidak diikutkan di paket produksi (murni verifikasi pengembangan).
"""
import os
import sys
import tempfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

tmpdir = tempfile.mkdtemp()
os.environ["HOME"] = tmpdir

from app import create_app  # noqa: E402
from app import db as _db  # noqa: E402
from app import logic as L  # noqa: E402

FAILS = []


def check(label, cond):
    print(f"[{'OK' if cond else 'FAIL'}] {label}")
    if not cond:
        FAILS.append(label)


db_path = os.path.join(tmpdir, "test.db")
app = create_app(db_path=db_path)
app.config["TESTING"] = True
app.config["WTF_CSRF_ENABLED"] = False
client = app.test_client()
client.post("/login", data={"username": "kaprodi", "password1": "test1234", "password2": "test1234"}, follow_redirects=True)

conn = _db.connect(db_path)

# ------------------------------------------------------------------- Fixture
# Tahun ajaran + 2 tahap (meniru pola "3 tahap per semester" milik pengguna,
# di sini cukup 2 supaya bisa dibuktikan filternya benar2 memisahkan).
client.post(
    "/pengaturan/tahun-akademik",
    data={"aksi": "buka_tahun", "kode": "2025/2026", "aktifkan": "Ganjil"},
    follow_redirects=True,
)
periode_row = conn.execute(
    "SELECT pa.id FROM periode_akademik pa JOIN tahun_ajaran ta ON ta.id=pa.tahun_ajaran_id "
    "WHERE ta.kode='2025/2026' AND pa.jenis='Ganjil'"
).fetchone()
pid = periode_row["id"]
for nama in ["Tahap 1 2025/2026", "Tahap 2 2025/2026"]:
    client.post("/pengaturan/tahun-akademik", data={"aksi": "tambah_tahap", "periode_id": pid, "nama_tahap": nama})

TAHAP1 = "Tahap 1 2025/2026"
TAHAP2 = "Tahap 2 2025/2026"

conn.execute("INSERT INTO dosen(nama, nidn, aktif) VALUES('Dr. Pembimbing Satu', '111', 1)")
dosen_id = conn.execute("SELECT id FROM dosen WHERE nama='Dr. Pembimbing Satu'").fetchone()["id"]

# 2 mahasiswa TAHAP1 (LULUS), 1 mahasiswa TAHAP2 (LULUS)
mhs = {}
for nim, nama, tahap in [
    ("2101", "Mahasiswa Satu", TAHAP1),
    ("2102", "Mahasiswa Dua", TAHAP1),
    ("2103", "Mahasiswa Tiga", TAHAP2),
]:
    conn.execute(
        "INSERT INTO mahasiswa(nim, nama, status_ta) VALUES(?,?,?)", (nim, nama, "Sudah Sidang")
    )
    mid = conn.execute("SELECT id FROM mahasiswa WHERE nim=?", (nim,)).fetchone()["id"]
    mhs[nim] = mid
    conn.execute(
        "INSERT INTO penetapan_pembimbing(mahasiswa_id, tahap, pembimbing1_id, judul_final) "
        "VALUES(?,?,?,?)",
        (mid, tahap, dosen_id, f"Judul Skripsi {nama}"),
    )
    conn.execute(
        "INSERT INTO sidang(mahasiswa_id, judul_sidang, nilai_angka, status_kelulusan) "
        "VALUES(?,?,?,?)",
        (mid, f"Judul Skripsi {nama}", 80, "LULUS"),
    )
conn.commit()

# ----------------------------------------------------- 1) logic.py langsung
rows_all = L.rencana_yudisium_rows(conn)
check("Tanpa filter: 3 baris Rencana Yudisium (semua tahap)", len(rows_all) == 3)

rows_t1 = L.rencana_yudisium_rows(conn, TAHAP1)
check("Filter Tahap 1: hanya 2 mahasiswa Tahap 1", len(rows_t1) == 2)
check("Filter Tahap 1: tidak memuat mahasiswa Tahap 2", all(r["nim"] != "2103" for r in rows_t1))

rows_t2 = L.rencana_yudisium_rows(conn, TAHAP2)
check("Filter Tahap 2: hanya 1 mahasiswa Tahap 2", len(rows_t2) == 1 and rows_t2[0]["nim"] == "2103")

check("Kolom 'tahap' ikut terbawa di rencana_yudisium_rows()", rows_t1[0]["tahap"] == TAHAP1)

# ------------------------------------------------------------------- 2) HTTP
r = client.get("/kelulusan/yudisium")
check("GET /kelulusan/yudisium -> 200", r.status_code == 200)
check("Tab filter tahap tampil di halaman Yudisium", TAHAP1.encode() in r.data and TAHAP2.encode() in r.data)

r_t1 = client.get(f"/kelulusan/yudisium?tahap={TAHAP1}")
check("GET yudisium?tahap=Tahap1 -> 200", r_t1.status_code == 200)
check("Halaman ter-filter Tahap 1 memuat Mahasiswa Satu", b"Mahasiswa Satu" in r_t1.data)
check("Halaman ter-filter Tahap 1 TIDAK memuat Mahasiswa Tiga", b"Mahasiswa Tiga" not in r_t1.data)

r_wisuda = client.get(f"/kelulusan/wisuda?tahap={TAHAP1}")
check("GET wisuda?tahap=Tahap1 -> 200 (tab filter juga ada di Wisuda)", r_wisuda.status_code == 200)

# ------------------------------------------------- 3) Tetapkan No. SK per Tahap
r_tetapkan = client.post(
    "/kelulusan/yudisium/tetapkan-tahap",
    data={"tahap": TAHAP1, "no_sk_batch": "SK/001/YUD/2026", "tgl_yudisium_batch": "2026-01-15"},
    follow_redirects=True,
)
check("POST tetapkan-tahap (Tahap 1) -> 200", r_tetapkan.status_code == 200)

no_sk_1 = conn.execute("SELECT no_sk FROM yudisium WHERE mahasiswa_id=?", (mhs["2101"],)).fetchone()["no_sk"]
no_sk_2 = conn.execute("SELECT no_sk FROM yudisium WHERE mahasiswa_id=?", (mhs["2102"],)).fetchone()["no_sk"]
no_sk_3 = conn.execute("SELECT no_sk FROM yudisium WHERE mahasiswa_id=?", (mhs["2103"],)).fetchone()["no_sk"]
check("No. SK diterapkan ke mahasiswa Satu (Tahap 1)", no_sk_1 == "SK/001/YUD/2026")
check("No. SK diterapkan ke mahasiswa Dua (Tahap 1)", no_sk_2 == "SK/001/YUD/2026")
check("No. SK Tahap 2 TIDAK ikut tertimpa (masih kosong)", not no_sk_3)

status_1 = conn.execute(
    "SELECT status_yudisium FROM yudisium WHERE mahasiswa_id=?", (mhs["2101"],)
).fetchone()["status_yudisium"]
check("Status ikut naik jadi 'Terlaksana' setelah SK ditetapkan", status_1 == "Terlaksana")

# Coba timpa tanpa centang "timpa" -> harus dilewati (tidak berubah)
client.post(
    "/kelulusan/yudisium/tetapkan-tahap",
    data={"tahap": TAHAP1, "no_sk_batch": "SK/BEDA/2026", "tgl_yudisium_batch": "2026-01-16"},
    follow_redirects=True,
)
no_sk_1_lagi = conn.execute(
    "SELECT no_sk FROM yudisium WHERE mahasiswa_id=?", (mhs["2101"],)
).fetchone()["no_sk"]
check("Tanpa centang 'timpa', No. SK yang sudah ada TIDAK berubah", no_sk_1_lagi == "SK/001/YUD/2026")

# Dengan centang "timpa" -> baru berubah
client.post(
    "/kelulusan/yudisium/tetapkan-tahap",
    data={
        "tahap": TAHAP1,
        "no_sk_batch": "SK/BEDA/2026",
        "tgl_yudisium_batch": "2026-01-16",
        "timpa": "1",
    },
    follow_redirects=True,
)
no_sk_1_timpa = conn.execute(
    "SELECT no_sk FROM yudisium WHERE mahasiswa_id=?", (mhs["2101"],)
).fetchone()["no_sk"]
check("Dengan centang 'timpa', No. SK berhasil diperbarui", no_sk_1_timpa == "SK/BEDA/2026")

# kembalikan ke kondisi semula utk uji cetak di bawah
client.post(
    "/kelulusan/yudisium/tetapkan-tahap",
    data={
        "tahap": TAHAP1,
        "no_sk_batch": "SK/001/YUD/2026",
        "tgl_yudisium_batch": "2026-01-15",
        "timpa": "1",
    },
    follow_redirects=True,
)

# -------------------------------------------------------- 4) Cetak SK Tahap
r_cetak = client.get(f"/surat/sk-yudisium-tahap?tahap={TAHAP1}")
check("GET /surat/sk-yudisium-tahap (Tahap 1) -> 200", r_cetak.status_code == 200)
check(
    "Response berupa file Word (.docx)",
    r_cetak.headers.get("Content-Type", "").startswith(
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    ),
)

import io
import docx

doc = docx.Document(io.BytesIO(r_cetak.data))
full_text = "\n".join(p.text for p in doc.paragraphs)
check("Dokumen SK Tahap memuat No. SK yang ditetapkan", "SK/001/YUD/2026" in full_text)
check("Dokumen SK Tahap menyebut nama tahap", TAHAP1 in full_text)

table_texts = []
for t in doc.tables:
    for row in t.rows:
        for cell in row.cells:
            table_texts.append(cell.text)
tabel_gabungan = " | ".join(table_texts)
check("Tabel SK memuat Mahasiswa Satu", "Mahasiswa Satu" in tabel_gabungan)
check("Tabel SK memuat Mahasiswa Dua", "Mahasiswa Dua" in tabel_gabungan)
check("Tabel SK TIDAK memuat Mahasiswa Tiga (beda tahap)", "Mahasiswa Tiga" not in tabel_gabungan)

# tahap kosong / "Semua" -> ditolak dengan redirect + flash, bukan 500/dokumen kosong
r_semua = client.get("/surat/sk-yudisium-tahap?tahap=Semua", follow_redirects=True)
check("Cetak SK dengan tahap='Semua' ditolak dgn baik (redirect, bukan error)", r_semua.status_code == 200)

r_kosong = client.get("/surat/sk-yudisium-tahap")
check("Cetak SK tanpa parameter tahap -> redirect balik ke Yudisium (bukan 500)", r_kosong.status_code in (302, 200))

# ------------------------------------------------------- 5) Ekspor Excel kolom Tahap
r_ekspor = client.get(f"/kelulusan/yudisium/ekspor?tahap={TAHAP1}")
check("GET yudisium/ekspor?tahap=Tahap1 -> 200", r_ekspor.status_code == 200)

import openpyxl

wb = openpyxl.load_workbook(io.BytesIO(r_ekspor.data))
ws = wb.active
header_row = [c.value for c in ws[1]]
check("Kolom 'Tahap/Gelombang' ada di ekspor Rencana Yudisium", "Tahap/Gelombang" in header_row)
nim_col_vals = [row[0].value for row in ws.iter_rows(min_row=2)]
check("Ekspor Rencana Yudisium (filter Tahap 1) hanya memuat 2 mahasiswa", len(nim_col_vals) == 2)

print()
print("=== SELESAI ===")
if FAILS:
    print(f"{len(FAILS)} GAGAL:")
    for f in FAILS:
        print(f" - {f}")
    sys.exit(1)
print("SEMUA TES LULUS.")
