# -*- coding: utf-8 -*-
"""routes/surat_umum.py — Modul 8: Generator Surat Umum (di luar Tugas Akhir)
+ Buku Agenda Surat Keluar.

Melengkapi `routes/surat.py` (yang khusus SK Pembimbing/Yudisium & Undangan
Seminar/Sidang, terikat pada data mahasiswa) dengan generator surat resmi
UMUM: Surat Tugas, Surat Keterangan, Surat Keputusan, Surat Undangan,
Surat Edaran, Nota Dinas, dst — isi surat diketik bebas oleh Kaprodi
(setiap jenis surat administratif punya kebutuhan isi yang berbeda-beda,
tidak realistis dibuatkan form kaku per jenis), sementara kop surat, nomor
otomatis, dan blok tanda tangan tetap dibuat otomatis dari Pengaturan
Identitas & Branding — konsisten dengan `_header()`/`_footer_ttd()` di
surat.py.

Setiap surat yang dibuat dicatat di tabel `surat_keluar` (Buku Agenda Surat
Keluar) sehingga nomor surat tidak pernah tabrakan (nomor urut dihitung
dari MAX nomor urut tahun berjalan +1 — bukan COUNT(*)+1, supaya tidak
pernah menurun/dipakai ulang walau ada baris yang dihapus di tengah
urutan) dan file .docx yang sudah digenerate bisa diunduh ulang kapan
saja tanpa perlu generate ulang.
"""

import io
import os
from datetime import date

from flask import (
    Blueprint,
    abort,
    current_app,
    flash,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
)

from app import constants as C
from app import datetools as dt
from app import db as _db
from app import error_utils as EH

bp = Blueprint("surat_umum", __name__, url_prefix="/surat-umum")


def _folder():
    folder = os.path.join(_db.home_dir(), "SistemSkripsi", "surat_keluar")
    os.makedirs(folder, exist_ok=True)
    return folder


def _nomor_otomatis(conn, jenis_surat, tgl):
    """Nomor urut surat = MAX(urut yang sudah pernah dipakai tahun berjalan) + 1
    -- BUKAN COUNT(*)+1.

    COUNT(*)+1 tabrakan begitu ada surat di tengah urutan yang dihapus:
    COUNT menurun sehingga urut berikutnya dihitung ulang dan bisa sama
    dengan nomor yang MASIH dipakai surat lain yang belum dihapus (lihat
    ADDENDUM_N1_REKAP-sibling bug report). MAX(...)+1 tidak pernah menurun
    walau ada penghapusan, jadi konsisten dengan klaim "nomor tidak
    dipakai ulang" di `hapus()` & docstring modul ini.

    Nomor urut diparse dari 3 digit pertama `nomor_surat` yang tersimpan
    (bukan dari kolom terpisah) supaya tetap benar walau kode jenis surat
    atau kode institusi (Pengaturan) berubah di antara pembuatan surat.
    """
    kode_jenis = C.JENIS_SURAT_UMUM.get(jenis_surat, "SRT")
    kode_institusi = _db.get_setting(conn, "kode_institusi_surat", "SIMPRODI")
    tahun = tgl.year
    rows = conn.execute(
        "SELECT nomor_surat FROM surat_keluar WHERE tanggal_surat LIKE ?", (f"{tahun}-%",)
    ).fetchall()
    urut_terpakai = 0
    for r in rows:
        depan = (r["nomor_surat"] or "").split("/", 1)[0]
        if depan.isdigit():
            urut_terpakai = max(urut_terpakai, int(depan))
    urut = urut_terpakai + 1
    romawi = dt.bulan_ke_romawi(tgl.month)
    return f"{urut:03d}/{kode_jenis}/{kode_institusi}/{romawi}/{tahun}"


def _bangun_docx(
    conn, jenis_surat, nomor, perihal, tujuan, tgl, isi, penandatangan, jabatan, tembusan,
    nip_nidn="",
):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    doc = docx.Document()

    # --- Kop surat, sama gaya dengan _header() di surat.py ---
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_db.get_setting(conn, "nama_institusi"))
    run.bold = True
    run.font.size = Pt(13)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(
        f"{_db.get_setting(conn, 'nama_fakultas')} | {_db.get_setting(conn, 'nama_prodi')}"
    ).font.size = Pt(10)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    alamat = _db.get_setting(conn, "alamat")
    telp = _db.get_setting(conn, "telp")
    email = _db.get_setting(conn, "email")
    p3.add_run(f"{alamat} · Telp {telp} · {email}").font.size = Pt(8.5)
    doc.add_paragraph("=" * 78).alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- Judul jenis surat ---
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(jenis_surat.upper())
    r.bold = True
    r.underline = True
    r.font.size = Pt(13)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.add_run(f"Nomor: {nomor}").font.size = Pt(11)
    doc.add_paragraph("")

    if tujuan:
        doc.add_paragraph(f"Kepada Yth.\n{tujuan}\ndi Tempat")
        doc.add_paragraph("")

    doc.add_paragraph(f"Perihal: {perihal}")
    doc.add_paragraph("")

    pembuka = C.PEMBUKA_SURAT_UMUM.get(jenis_surat, "")
    if pembuka:
        doc.add_paragraph(pembuka)
        doc.add_paragraph("")

    for para in (isi or "").split("\n\n"):
        para = para.strip()
        if para:
            doc.add_paragraph(para)

    doc.add_paragraph("")
    doc.add_paragraph(
        "Demikian surat ini dibuat untuk dapat dipergunakan sebagaimana mestinya. "
        "Atas perhatian dan kerja samanya, kami ucapkan terima kasih."
    )
    doc.add_paragraph("")

    # --- Blok tanda tangan ---
    ttd = doc.add_paragraph()
    ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    ttd.add_run(f"{dt.format_tanggal(tgl)}\n{jabatan or 'Ketua Program Studi'},")
    for _ in range(3):
        doc.add_paragraph("").alignment = WD_ALIGN_PARAGRAPH.RIGHT
    nama_ttd = doc.add_paragraph()
    nama_ttd.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = nama_ttd.add_run(penandatangan or "(...........................................)")
    r2.underline = True
    if nip_nidn:
        nip_p = doc.add_paragraph()
        nip_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        nip_p.add_run(f"NIP/NIDN. {nip_nidn}")

    if tembusan:
        doc.add_paragraph("")
        doc.add_paragraph("Tembusan:")
        for i, t in enumerate([x.strip() for x in tembusan.splitlines() if x.strip()], start=1):
            doc.add_paragraph(f"{i}. {t}")

    return doc


@bp.route("/")
def index():
    conn = current_app.get_db()
    jenis_filter = request.args.get("jenis", "")
    cari = request.args.get("cari", "").strip()

    q = "SELECT * FROM surat_keluar WHERE 1=1"
    params = []
    if jenis_filter:
        q += " AND jenis_surat=?"
        params.append(jenis_filter)
    if cari:
        # Audit UI/UX — modul ini sebelumnya tidak punya pencarian sama
        # sekali (beda dari Document Center & modul list+form lain yang
        # semuanya punya kotak "Cari"), padahal Buku Agenda bisa berisi
        # ratusan surat dan chip jenis saja tidak cukup untuk menemukan
        # satu surat tertentu.
        q += " AND (nomor_surat LIKE ? OR perihal LIKE ? OR tujuan LIKE ?)"
        like = f"%{cari}%"
        params += [like, like, like]
    q += " ORDER BY id DESC"
    rows = conn.execute(q, params).fetchall()

    jumlah_per_jenis = {
        r["jenis_surat"]: r["c"]
        for r in conn.execute(
            "SELECT jenis_surat, COUNT(*) c FROM surat_keluar GROUP BY jenis_surat"
        ).fetchall()
    }
    total_keseluruhan = sum(jumlah_per_jenis.values())
    tahun_ini = date.today().year
    total_tahun_ini = conn.execute(
        "SELECT COUNT(*) c FROM surat_keluar WHERE tanggal_surat LIKE ?", (f"{tahun_ini}-%",)
    ).fetchone()["c"]
    bulan_ini = conn.execute(
        "SELECT COUNT(*) c FROM surat_keluar WHERE strftime('%Y-%m', dibuat_pada)=strftime('%Y-%m','now','localtime')"
    ).fetchone()["c"]
    jenis_terbanyak = (
        max(jumlah_per_jenis, key=jumlah_per_jenis.get) if jumlah_per_jenis else None
    )

    nama_default = _db.get_setting(conn, "nama_penandatangan_default", "")
    jabatan_default = _db.get_setting(conn, "jabatan_penandatangan_default", "Ketua Program Studi")

    return render_template(
        "surat_umum.html",
        rows=rows,
        jenis_list=C.JENIS_SURAT_UMUM_LIST,
        jenis_filter=jenis_filter,
        cari=cari,
        today=date.today().isoformat(),
        nama_default=nama_default,
        jabatan_default=jabatan_default,
        total_keseluruhan=total_keseluruhan,
        total_tahun_ini=total_tahun_ini,
        tahun_ini=tahun_ini,
        bulan_ini=bulan_ini,
        jenis_terbanyak=jenis_terbanyak,
    )


@bp.route("/buat", methods=["POST"])
def buat():
    conn = current_app.get_db()
    f = request.form
    jenis_surat = f.get("jenis_surat", "")
    perihal = f.get("perihal", "").strip()
    if jenis_surat not in C.JENIS_SURAT_UMUM or not perihal:
        flash("Jenis Surat dan Perihal wajib diisi.", "error")
        return redirect(url_for("surat_umum.index"))

    tgl_text = f.get("tanggal_surat", "").strip()
    tgl = dt.parse_tanggal(tgl_text) or date.today()
    tujuan = f.get("tujuan", "").strip()
    isi = f.get("isi", "").strip()
    penandatangan = f.get("penandatangan", "").strip() or _db.get_setting(
        conn, "nama_penandatangan_default", ""
    )
    jabatan = f.get("jabatan_penandatangan", "").strip() or _db.get_setting(
        conn, "jabatan_penandatangan_default", "Ketua Program Studi"
    )
    # Fase Pejabat Struktural — NIP/NIDN ikut dicetak di bawah nama
    # penandatangan kalau tersedia (disinkron otomatis dari Pengaturan >
    # Pejabat Struktural lewat "Jadikan Default Penandatangan"). Hanya
    # dipakai saat penandatangan TIDAK diubah manual di form, supaya NIP
    # pejabat lama tidak salah tempel ke nama yang diketik manual.
    nip_nidn = ""
    if not f.get("penandatangan", "").strip():
        nip_nidn = _db.get_setting(conn, "nip_nidn_penandatangan_default", "")
    tembusan = f.get("tembusan", "").strip()

    # Audit Lanjutan 11 — SEBELUM perbaikan ini, seluruh alur di bawah
    # (nomor otomatis + generate .docx + tulis file ke disk + INSERT ke
    # Buku Agenda + kirim unduhan) TIDAK dibungkus try/except sama sekali
    # — persis bug yang sama yang ditemukan & ditambal di
    # `routes/surat.py::buat()` pada Audit Lanjutan 10. Diverifikasi
    # nyata: galat tak terduga (mis. `get_setting` gagal baca Pengaturan)
    # bocor jadi 500 Flask mentah alih-alih flash pesan ramah.
    try:
        nomor = _nomor_otomatis(conn, jenis_surat, tgl)
        doc = _bangun_docx(
            conn, jenis_surat, nomor, perihal, tujuan, tgl, isi, penandatangan, jabatan, tembusan,
            nip_nidn=nip_nidn,
        )

        aman_nomor = nomor.replace("/", "-")
        filename = f"{aman_nomor}_{jenis_surat.replace(' ', '_')}.docx"
        dest = os.path.join(_folder(), filename)
        doc.save(dest)

        conn.execute(
            "INSERT INTO surat_keluar(nomor_surat, jenis_surat, perihal, tujuan, tanggal_surat, "
            "isi_ringkas, penandatangan, jabatan_penandatangan, tembusan, lokasi_file) "
            "VALUES(?,?,?,?,?,?,?,?,?,?)",
            (
                nomor,
                jenis_surat,
                perihal,
                tujuan,
                tgl.isoformat(),
                isi,
                penandatangan,
                jabatan,
                tembusan,
                dest,
            ),
        )
        conn.commit()
        _db.log(conn, "Buat Surat Umum", f"{jenis_surat} — {nomor}")
        flash(f"Surat {nomor} berhasil dibuat & tercatat di Buku Agenda.", "ok")

        buf = io.BytesIO()
        doc.save(buf)
        buf.seek(0)
        return send_file(
            buf,
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
    except Exception as e:
        EH.flash_gagal_simpan(e, f"Gagal membuat dokumen {jenis_surat}")
        return redirect(url_for("surat_umum.index"))


@bp.route("/<int:sid>/unduh")
def unduh(sid):
    conn = current_app.get_db()
    row = conn.execute("SELECT * FROM surat_keluar WHERE id=?", (sid,)).fetchone()
    if not row or not row["lokasi_file"] or not os.path.exists(row["lokasi_file"]):
        abort(404)
    nama = os.path.basename(row["lokasi_file"])
    return send_file(row["lokasi_file"], as_attachment=True, download_name=nama)


@bp.route("/<int:sid>/hapus", methods=["POST"])
def hapus(sid):
    conn = current_app.get_db()
    row = conn.execute("SELECT * FROM surat_keluar WHERE id=?", (sid,)).fetchone()
    if row:
        if row["lokasi_file"] and os.path.exists(row["lokasi_file"]):
            try:
                os.remove(row["lokasi_file"])
            except OSError:
                pass
        conn.execute("DELETE FROM surat_keluar WHERE id=?", (sid,))
        conn.commit()
        _db.log(conn, "Hapus Surat Umum (agenda)", row["nomor_surat"] or str(sid))
    flash("Catatan surat dihapus dari Buku Agenda (nomor tidak dipakai ulang).", "ok")
    return redirect(url_for("surat_umum.index"))
