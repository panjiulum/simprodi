# -*- coding: utf-8 -*-
"""routes/panduan.py — Panduan Penggunaan Aplikasi (User Guide).

Modul referensi/dokumentasi in-app: menjelaskan cara memakai setiap modul
SIMPRODI secara lengkap (fungsi, langkah pemakaian, dan tips) tanpa
menyentuh logika/skema modul lain. Murni read-only — tidak ada tabel baru,
tidak ada endpoint yang menulis data.

Konten diorganisir mengikuti struktur grup sidebar (lihat
docs/RESTRUKTURISASI_SIDEBAR.md) supaya urutan panduan langsung sejalan
dengan urutan menu yang dilihat pengguna, dan disajikan sebagai accordion
per grup (pola UI yang sama dengan sidebar "expand" hasil restrukturisasi)
supaya halaman tidak jadi satu scroll raksasa.

Catatan lingkup: menu di grup Pengaturan yang masih berstatus roadmap
(placeholder `routes/roadmap.py`) tetap disebutkan di sini secara singkat
(supaya panduan tetap 1:1 dengan sidebar), tapi ditandai jelas sebagai
"belum tersedia" alih-alih dijelaskan seolah-olah sudah berfungsi.
"""

import io
import os
from datetime import date

from flask import Blueprint, current_app, render_template, send_file

from app import constants as C
from app import db as _db

bp = Blueprint("panduan", __name__, url_prefix="/panduan")

# Setiap grup: (label_grup, ringkasan_grup, [modul, ...])
# Setiap modul: dict dengan:
#   judul       - nama modul (disamakan dengan label di sidebar)
#   endpoint    - endpoint Flask yang dituju menu ini (untuk tautan "Buka modul")
#   ringkasan   - 1-2 kalimat: modul ini untuk apa
#   langkah     - list langkah pemakaian, urut
#   tips        - list catatan/tips tambahan (boleh kosong)
PANDUAN_GROUPS = [
    (
        "🎓 Akademik",
        "Operasional akademik reguler yang dipakai rutin tiap minggu: "
        "kurikulum & OBE, jadwal kelas, input nilai, dan kalender akademik.",
        [
            dict(
                judul="Kurikulum & OBE",
                endpoint="kurikulum.index",
                ringkasan=(
                    "Mengelola Capaian Pembelajaran Lulusan (CPL), struktur mata "
                    "kuliah, pemetaan CPMK-ke-CPL, dan status RPS per mata kuliah — "
                    "fondasi data untuk seluruh rantai keterlacakan OBE di SIMPRODI."
                ),
                langkah=[
                    "Buka tab Dashboard OBE untuk melihat ringkasan jumlah CPL, mata kuliah, dan status RPS.",
                    "Di tab Struktur Kurikulum, tambahkan/lengkapi data mata kuliah (kode, nama, SKS, semester).",
                    "Isi CPL di 4 kategori SN-DIKTI: Sikap, Pengetahuan, Keterampilan Umum, Keterampilan Khusus.",
                    "Di tab Pemetaan CPL-CPMK, tautkan setiap CPMK mata kuliah ke satu atau lebih CPL — pemetaan ini yang jadi dasar perhitungan capaian OBE di Modul Nilai.",
                    "Di tab RPS & Perangkat, tandai status kelengkapan RPS per mata kuliah.",
                ],
                tips=[
                    "Lengkapi Kurikulum & pemetaan CPMK-CPL sebelum mulai memakai Modul Jadwal dan Nilai — kedua modul itu bergantung pada mata kuliah kurikulum aktif.",
                    "Kategori CPL sengaja mengikuti 4 unsur SN-DIKTI supaya bisa langsung dipakai sebagai matriks borang akreditasi LAMEMBA.",
                ],
            ),
            dict(
                judul="Jadwal Kelas & BAP",
                endpoint="jadwal.index",
                ringkasan=(
                    "Mengelola jadwal kelas per mata kuliah pada tahun akademik/"
                    "semester berjalan, dan mencatat log Berita Acara Perkuliahan "
                    "(BAP) per pertemuan."
                ),
                langkah=[
                    "Di tab Jadwal Kelas, buat kelas baru untuk mata kuliah dari kurikulum aktif (pilih dosen pengampu, hari/jam, ruangan).",
                    "Di tab BAP, catat log tiap pertemuan: materi yang diajarkan, Sub-CPMK yang diacu, dan kehadiran.",
                    "Realisasi jumlah pertemuan dihitung otomatis dari log BAP yang sudah diisi — tidak perlu diisi manual.",
                ],
                tips=[
                    "Kelas yang dibuat di sini otomatis jadi opsi 'peserta kelas' di Modul Nilai.",
                    "Sistem akan menandai bentrok jadwal (ruangan/dosen) memakai logika yang sama dengan Modul Seminar/Sidang.",
                ],
            ),
            dict(
                judul="Nilai & OBE Assessment",
                endpoint="nilai.index",
                ringkasan=(
                    "Assessment Engine OBE lengkap: dari KRS per kelas, input nilai "
                    "per CPMK, sampai rekap capaian CPL individu maupun tingkat "
                    "program studi."
                ),
                langkah=[
                    "Tab Peserta Kelas: daftarkan mahasiswa ke satu jadwal kelas (setara KRS).",
                    "Tab Input Nilai: isi nilai per CPMK untuk tiap mahasiswa terdaftar, lalu nilai akhir kelas.",
                    "Tab Capaian CPL Individu: lihat rekap capaian OBE 1 mahasiswa lintas seluruh mata kuliah yang sudah dinilai.",
                    "Tab Capaian CPL Program: lihat rekap capaian OBE tingkat program studi (dipakai juga oleh Modul Siklus CQI).",
                ],
                tips=[
                    "Urutan wajib: Kurikulum (CPL/CPMK) → Jadwal (kelas) → Nilai. Kalau capaian CPL tampak kosong, cek dulu apakah pemetaan CPMK-CPL di Kurikulum sudah lengkap.",
                ],
            ),
            dict(
                judul="Kalender Akademik",
                endpoint="kalender.index",
                ringkasan=(
                    "Agenda akademik prodi (ujian, rapat, libur, deadline, kegiatan) "
                    "dalam tampilan kalender bulanan."
                ),
                langkah=[
                    "Buka bulan yang ingin dilihat memakai navigasi bulan pada kalender.",
                    "Tambahkan agenda baru lewat form (judul, tanggal, kategori/jenis agenda).",
                    "Klik salah satu tanggal berisi agenda untuk melihat/mengubah/menghapus detailnya.",
                ],
                tips=[
                    "Ikon lonceng 🔔 di pojok kanan atas topbar (semua halaman) adalah pintasan langsung ke Kalender & Reminder.",
                ],
            ),
        ],
    ),
    (
        "👤 Mahasiswa",
        "Siklus hidup mahasiswa dari aktif kuliah reguler, jalur non-reguler "
        "(Semester Pendek, RPL), sampai tracer study setelah lulus.",
        [
            dict(
                judul="Data Mahasiswa",
                endpoint="mahasiswa.list_view",
                ringkasan="Data induk seluruh mahasiswa (identitas, NIM, skema Reguler/RPL).",
                langkah=[
                    "Gunakan kolom pencarian untuk mencari mahasiswa berdasarkan nama/NIM.",
                    "Gunakan filter Skema untuk menyaring mahasiswa Reguler atau RPL.",
                    "Klik Tambah untuk mendaftarkan mahasiswa baru, atau ikon edit pada baris untuk mengubah data.",
                    "Gunakan tombol ekspor untuk mengunduh data mahasiswa ke Excel.",
                ],
                tips=[
                    "Data mahasiswa di sini dipakai lintas hampir semua modul lain (Tugas Akhir, Nilai, Semester Pendek, RPL, Tracer Study) — pastikan NIM dan nama benar sejak awal.",
                ],
            ),
            dict(
                judul="Semester Pendek",
                endpoint="sp.index",
                ringkasan=(
                    "Pengelolaan Semester Pendek (SP) dari konfigurasi periode, "
                    "penawaran kelas, pendaftaran, sampai penilaian."
                ),
                langkah=[
                    "Tab Periode: buat/konfigurasi periode SP (timeline, aturan akademik, biaya).",
                    "Tab Kelas: buka penawaran kelas SP untuk mata kuliah dari kurikulum aktif — status kuota (Dibuka/Kurang Kuota/Penuh) terhitung otomatis.",
                    "Tab Peserta: proses pendaftaran mahasiswa (alasan mengulang & nilai sebelumnya wajib dicatat), lalu setujui/tolak pendaftaran secara manual.",
                    "Lanjutkan ke tab Pelaksanaan dan Penilaian sesuai progres periode SP berjalan.",
                ],
                tips=[
                    "Persetujuan peserta SP sengaja dibuat manual (bukan otomatis) supaya Kaprodi tetap memverifikasi syarat mengulang mata kuliah bernilai di bawah C.",
                ],
            ),
            dict(
                judul="RPL",
                endpoint="rpl.index",
                ringkasan=(
                    "Rekognisi Pembelajaran Lampau: alur pendaftar, asesmen "
                    "konversi SKS, sampai dokumen pendukung per pendaftar."
                ),
                langkah=[
                    "Tab Pendaftar: catat data pendaftar RPL dan ubah status tahap asesmen (Verifikasi Berkas → Asesmen Portofolio → Disetujui/Ditolak).",
                    "Pilih satu pendaftar, lalu buka tab Asesmen untuk mengisi konversi SKS per mata kuliah — total SKS diakui terhitung otomatis.",
                    "Tab Dokumen: unggah/unduh/hapus berkas pendukung untuk pendaftar terpilih.",
                ],
                tips=[
                    "Tab Asesmen dan Dokumen selalu terikat ke satu pendaftar yang sedang dipilih dari tab Pendaftar.",
                ],
            ),
            dict(
                judul="Tracer Study",
                endpoint="kelulusan.tracer_list",
                ringkasan="Pelacakan alumni pasca-lulus (status karier, umpan balik) untuk kebutuhan akreditasi.",
                langkah=[
                    "Lihat daftar alumni beserta status pengisian tracer study.",
                    "Klik salah satu alumni untuk melengkapi/memperbarui data tracer study-nya.",
                ],
                tips=[],
            ),
        ],
    ),
    (
        "📚 Tugas Akhir",
        "Alur kerja linear 6 tahap Tugas Akhir mahasiswa: Pengajuan Judul → "
        "Penetapan Pembimbing → Seminar Proposal → Sidang Skripsi → "
        "Rencana Yudisium → Wisuda.",
        [
            dict(
                judul="Pengajuan Judul",
                endpoint="akademik.pengajuan_list",
                ringkasan="Pengajuan dan review judul Tugas Akhir mahasiswa — tahap pertama alur TA.",
                langkah=[
                    "Tambahkan pengajuan judul baru: pilih mahasiswa, isi judul yang diajukan.",
                    "Ubah status pengajuan (mis. Diajukan/Disetujui/Ditolak) sesuai hasil review Kaprodi.",
                ],
                tips=[
                    "Mahasiswa dengan judul berstatus disetujui baru bisa lanjut ke tahap Penetapan Pembimbing."
                ],
            ),
            dict(
                judul="Penetapan Pembimbing",
                endpoint="akademik.penetapan_list",
                ringkasan="Penetapan dosen pembimbing (utama/pendamping) untuk mahasiswa yang judulnya sudah disetujui.",
                langkah=[
                    "Pilih mahasiswa yang judulnya sudah disetujui.",
                    "Tetapkan dosen pembimbing dari daftar dosen aktif.",
                    "Simpan — data ini yang jadi dasar penerbitan SK Pembimbing di Modul Cetak Surat Tugas Akhir.",
                ],
                tips=[],
            ),
            dict(
                judul="Seminar Proposal",
                endpoint="pelaksanaan.seminar_list",
                ringkasan="Penjadwalan dan pencatatan hasil Seminar Proposal, lengkap dengan deteksi bentrok jadwal.",
                langkah=[
                    "Pilih mahasiswa yang sudah punya SK Pembimbing.",
                    "Jadwalkan seminar: tanggal, jam, ruangan, dan dosen penguji — sistem otomatis memeriksa bentrok jadwal ruangan/dosen dan memberi konfirmasi bila terdeteksi.",
                    "Setelah seminar berlangsung, catat hasil/status kelulusan seminar.",
                ],
                tips=[
                    "Konfirmasi bentrok jadwal memakai logika yang sama dengan Modul Sidang Skripsi dan Jadwal Kelas."
                ],
            ),
            dict(
                judul="Sidang Skripsi",
                endpoint="pelaksanaan.sidang_list",
                ringkasan="Penjadwalan dan pencatatan hasil Sidang Skripsi, dengan mekanisme deteksi bentrok yang sama seperti Seminar Proposal.",
                langkah=[
                    "Pilih mahasiswa yang sudah lulus tahap Seminar Proposal.",
                    "Jadwalkan sidang beserta dosen penguji; sistem memvalidasi bentrok jadwal.",
                    "Catat hasil sidang — status ini menjadi basis sinkronisasi otomatis ke Rencana Yudisium.",
                ],
                tips=[],
            ),
            dict(
                judul="Rencana Yudisium",
                endpoint="kelulusan.yudisium_list",
                ringkasan="Daftar mahasiswa siap yudisium — auto-tersinkron dari status kelulusan Sidang Skripsi.",
                langkah=[
                    "Buka halaman ini — daftar mahasiswa yang lulus sidang otomatis muncul (tidak perlu ditambah manual).",
                    "Lengkapi kolom tambahan yang masih perlu diisi manual (mis. tanggal/nomor SK Yudisium).",
                ],
                tips=[
                    "Kalau mahasiswa yang sudah lulus sidang belum muncul di sini, cek kembali status hasil sidangnya di Modul Sidang Skripsi."
                ],
            ),
            dict(
                judul="Wisuda",
                endpoint="kelulusan.wisuda_list",
                ringkasan="Daftar mahasiswa yang menempuh wisuda — auto-tersinkron dari data yudisium, tahap terakhir alur Tugas Akhir.",
                langkah=[
                    "Buka halaman ini — daftar auto-terisi dari mahasiswa yang sudah diyudisium.",
                    "Lengkapi kolom terkait periode/tanggal wisuda sesuai kebutuhan.",
                ],
                tips=[],
            ),
        ],
    ),
    (
        "🧑‍🏫 SDM",
        "Data dosen, kinerja SDM, dan aset fisik (ruangan) yang menunjang operasional prodi.",
        [
            dict(
                judul="Data Dosen",
                endpoint="dosen.list_view",
                ringkasan="Data induk seluruh dosen (identitas, NIDN, status aktif).",
                langkah=[
                    "Cari dosen memakai kolom pencarian (nama/NIDN).",
                    "Tambah/ubah data dosen lewat form CRUD standar.",
                ],
                tips=[
                    "Hanya dosen berstatus aktif yang muncul sebagai opsi pembimbing/penguji di Modul Tugas Akhir."
                ],
            ),
            dict(
                judul="SDM & Kinerja Dosen",
                endpoint="sdm.index",
                ringkasan=(
                    "Rekam jejak kinerja per dosen: aktivitas pendidikan, penelitian, "
                    "PKM, penunjang, luaran, peran akademik, timeline karier, dan "
                    "target kinerja."
                ),
                langkah=[
                    "Pilih satu dosen dari daftar.",
                    "Pilih tab sesuai jenis data yang ingin dicatat (Pendidikan/Penelitian/PKM/Penunjang/Luaran/Peran Akademik/Timeline Karier/Target Kinerja).",
                    "Tambahkan entri log baru sesuai jenis aktivitas tersebut.",
                    "Lihat Dashboard Kesiapan & realisasi target — dihitung ulang otomatis setiap halaman dibuka, jadi selalu mencerminkan data terbaru.",
                ],
                tips=[
                    "Data aktivitas Penelitian/PKM/Luaran yang diinput di sini adalah SATU-SATUNYA sumber data — Modul Tri Dharma hanya membaca & merekapnya lintas dosen, jadi jangan menginput ulang di tempat lain.",
                ],
            ),
            dict(
                judul="Aset & Ruangan",
                endpoint="ruangan.list_view",
                ringkasan="Data ruangan yang dipakai untuk deteksi bentrok jadwal (kelas, seminar, sidang).",
                langkah=[
                    "Tambah/ubah data ruangan (nama, kapasitas, lokasi).",
                ],
                tips=[
                    "Ruangan yang terdaftar di sini otomatis jadi opsi pemilihan lokasi di Modul Jadwal, Seminar, dan Sidang."
                ],
            ),
        ],
    ),
    (
        "🧪 Tri Dharma",
        "Rekap tingkat program studi atas aktivitas Penelitian, PKM, dan "
        "Publikasi/HKI seluruh dosen — datanya sama dengan yang diinput di SDM.",
        [
            dict(
                judul="Penelitian & PKM",
                endpoint="tridharma.index",
                ringkasan="Rekap & filter aktivitas Penelitian dan PKM lintas seluruh dosen dalam satu tampilan tingkat program studi.",
                langkah=[
                    "Buka tab Penelitian & PKM, gunakan filter jenis (Penelitian/PKM) untuk menyaring tampilan.",
                    "Gunakan filter tahun/periode bila tersedia untuk melihat tren dari waktu ke waktu.",
                ],
                tips=[
                    "Untuk MENAMBAH data aktivitas baru, input dilakukan di Modul SDM & Kinerja Dosen (per dosen) — modul ini murni rekap read-only lintas dosen."
                ],
            ),
            dict(
                judul="Publikasi & HKI",
                endpoint="tridharma.index",
                ringkasan="Rekap luaran akademik (publikasi, HKI, buku, prosiding) seluruh dosen dalam satu tampilan program studi.",
                langkah=[
                    "Buka tab Luaran Akademik untuk melihat rekap seluruh luaran dosen.",
                ],
                tips=[
                    "Sama seperti Penelitian & PKM, input data luaran dilakukan lewat Modul SDM & Kinerja Dosen."
                ],
            ),
        ],
    ),
    (
        "🤝 Kerja Sama",
        "Data mitra kerja sama eksternal beserta program, dokumen, dan evaluasinya.",
        [
            dict(
                judul="Kerja Sama & Mitra",
                endpoint="kerjasama.index",
                ringkasan="Executive dashboard mitra, dokumen MoU, program/implementasi, dan evaluasi kerja sama.",
                langkah=[
                    "Tab Mitra & Dokumen: tambahkan data mitra dan unggah dokumen MoU terkait.",
                    "Tab Program & Implementasi: catat program kerja sama konkret — bisa ditautkan opsional ke PIC dosen atau ke aktivitas Penelitian/PKM terkait.",
                    "Tab Evaluasi & Luaran: isi skor kepuasan mitra per program — indeks kepuasan pada dashboard dihitung dari skor riil ini, bukan angka statis.",
                ],
                tips=[],
            ),
        ],
    ),
    (
        "📊 Mutu & Analytics",
        "Satu rumah untuk seluruh siklus penjaminan mutu (CQI, AMI) dan seluruh "
        "laporan/rekap analitis program studi.",
        [
            dict(
                judul="Siklus CQI",
                endpoint="cqi.index",
                ringkasan="Continuous Quality Improvement berbasis capaian CPL — siklus PDCA tingkat CPL/OBE.",
                langkah=[
                    "Tab Gap Analysis: bandingkan capaian CPL program saat ini dengan target.",
                    "Untuk CPL dengan gap, klik tombol buka siklus CQI — sistem membekukan (snapshot) capaian saat itu supaya rencana tindak lanjut tidak berubah kalau ada nilai baru masuk kemudian.",
                    "Tab Siklus CQI: kelola rencana tindak lanjut (Plan-Do-Check-Act) untuk tiap siklus yang sudah dibuka.",
                ],
                tips=[
                    "Untuk mengevaluasi ulang, buka siklus BARU di tahun akademik berikutnya — jangan menimpa siklus yang sudah berjalan."
                ],
            ),
            dict(
                judul="IKU",
                endpoint="mutu.index",
                ringkasan="Pemantauan 8 Indikator Kinerja Utama (IKU) Kemendikbudristek untuk program studi.",
                langkah=["Buka tab IKU untuk melihat/memperbarui capaian tiap indikator."],
                tips=[],
            ),
            dict(
                judul="Akreditasi",
                endpoint="mutu.index",
                ringkasan="Pemantauan kesiapan borang berbasis 9 Kriteria LAMEMBA.",
                langkah=["Buka tab Akreditasi untuk melihat status kelengkapan tiap kriteria."],
                tips=[],
            ),
            dict(
                judul="Audit & QA",
                endpoint="mutu.index",
                ringkasan="Audit Mutu Internal (AMI/SPMI) — siklus PDCA tingkat standar mutu — dan pemindai kelengkapan data lintas modul.",
                langkah=[
                    "Buka tab Audit & QA untuk melihat/mengelola siklus AMI.",
                    "Gunakan pemindai kelengkapan data untuk melihat modul mana yang datanya belum lengkap.",
                ],
                tips=[
                    "AMI berbeda level dengan Siklus CQI: CQI mengevaluasi capaian CPL/OBE, AMI mengevaluasi kepatuhan terhadap standar SPMI."
                ],
            ),
            dict(
                judul="Rekap Pembimbing",
                endpoint="rekap.pembimbing",
                ringkasan="Rekap beban bimbingan Tugas Akhir per dosen pembimbing.",
                langkah=["Gunakan filter tahap/periode, lalu ekspor ke Excel bila diperlukan."],
                tips=[],
            ),
            dict(
                judul="Rekap Status",
                endpoint="rekap.status",
                ringkasan="Rekap status mahasiswa pada tiap tahap alur Tugas Akhir.",
                langkah=[
                    "Gunakan filter tahap untuk menyaring, lalu ekspor ke Excel bila diperlukan."
                ],
                tips=[],
            ),
            dict(
                judul="Rasio Beban Dosen",
                endpoint="rekap.rasio_dosen",
                ringkasan="Rasio jumlah bimbingan/beban terhadap jumlah dosen aktif.",
                langkah=["Lihat rekap rasio per dosen; ekspor ke Excel bila diperlukan."],
                tips=[],
            ),
            dict(
                judul="Rekap Kinerja Dosen (SDM)",
                endpoint="rekap.kinerja_dosen",
                ringkasan="Rekap jumlah kegiatan per kategori Tri Dharma (Pendidikan/Penelitian/PKM/Penunjang/Luaran/Peran Akademik/Timeline) + kesiapan BKD/SISTER per dosen, dari data Modul SDM.",
                langkah=[
                    "Gunakan filter Homebase/Semua Dosen dan Tahun Akademik.",
                    "Klik nama dosen untuk membuka detail lengkapnya di Modul SDM.",
                    "Ekspor ke Excel bila diperlukan untuk laporan BKD/SISTER/borang akreditasi.",
                ],
                tips=[
                    "Angka kesiapan BKD/SISTER di sini dihitung dengan rumus yang sama persis dengan Dashboard Modul SDM — selalu konsisten."
                ],
            ),
            dict(
                judul="Rekap Program Kerja",
                endpoint="rekap.program_kerja",
                ringkasan="Rekap realisasi Program Kerja tahunan per bidang, beserta detail tiap program dan jumlah kegiatan yang sudah Selesai.",
                langkah=[
                    "Gunakan filter Tahun Akademik untuk menyaring.",
                    "Lihat ringkasan per bidang di bagian atas, detail tiap program di bagian bawah.",
                    "Ekspor ke Excel (2 bagian: Detail Program + Ringkasan Bidang) bila diperlukan.",
                ],
                tips=[
                    "Persentase realisasi dihitung dengan rumus yang sama persis dengan halaman Program Kerja — selalu konsisten."
                ],
            ),
            dict(
                judul="Statistik",
                endpoint="rekap.statistik",
                ringkasan="Statistik agregat program studi (jumlah mahasiswa, kelulusan, dsb).",
                langkah=["Lihat ringkasan statistik; ekspor ke Excel bila diperlukan."],
                tips=[],
            ),
        ],
    ),
    (
        "📁 Operasional",
        "Alat bantu administratif yang dipakai lintas semua modul lain: program kerja, arsip dokumen, dan generator surat.",
        [
            dict(
                judul="Program Kerja",
                endpoint="kegiatan.index",
                ringkasan="Rencana program kerja tahunan prodi beserta realisasi kegiatan konkretnya, dalam satu halaman 2 tab.",
                langkah=[
                    "Tab Program Kerja: buat rencana program kerja tahunan.",
                    "Tab Kegiatan: catat kegiatan konkret yang terhubung ke satu program kerja, lalu tandai statusnya (mis. Selesai).",
                ],
                tips=[
                    "Persentase realisasi program kerja dihitung otomatis dari rasio kegiatan berstatus Selesai terhadap total kegiatan yang terhubung — tidak perlu dihitung manual."
                ],
            ),
            dict(
                judul="Document Center",
                endpoint="dokumen.index",
                ringkasan="Arsip dokumen prodi (SK, MoU, kurikulum, akreditasi, dll) — file fisik tersimpan di komputer, tautan & metadatanya di sini.",
                langkah=[
                    "Klik Unggah untuk menambahkan dokumen baru (pilih file, isi metadata seperti judul/kategori).",
                    "Gunakan tombol unduh pada baris dokumen untuk membuka kembali filenya.",
                    "Hapus dokumen yang sudah tidak relevan lewat tombol hapus (akan diminta konfirmasi).",
                ],
                tips=[],
            ),
            dict(
                judul="RKP Seminar",
                endpoint="rekap.rkp_seminar",
                ringkasan="Rekapitulasi kegiatan/biaya Seminar Proposal.",
                langkah=["Lihat rekap; ekspor ke Excel bila diperlukan."],
                tips=[],
            ),
            dict(
                judul="RKP Sidang",
                endpoint="rekap.rkp_sidang",
                ringkasan="Rekapitulasi kegiatan/biaya Sidang Skripsi.",
                langkah=["Lihat rekap; ekspor ke Excel bila diperlukan."],
                tips=[],
            ),
            dict(
                judul="Cetak Surat Tugas Akhir",
                endpoint="surat.index",
                ringkasan="Generator dokumen Word (.docx) siap cetak khusus alur Tugas Akhir: SK Pembimbing, SK Yudisium, Undangan Seminar/Sidang — terisi otomatis dari data yang sudah tersimpan.",
                langkah=[
                    "Pilih jenis surat (SK Pembimbing/SK Yudisium/Undangan Seminar/Undangan Sidang).",
                    "Pilih mahasiswa/data terkait — sistem otomatis mengisi field surat dari data yang sudah ada di modul lain.",
                    "Klik cetak/unduh untuk mendapatkan file .docx.",
                ],
                tips=[
                    "Untuk surat DI LUAR Tugas Akhir (Surat Tugas, Surat Keterangan, dll), gunakan Generator Surat Umum, bukan modul ini."
                ],
            ),
            dict(
                judul="Generator Surat Umum",
                endpoint="surat_umum.index",
                ringkasan="Generator surat resmi umum di luar Tugas Akhir (Surat Tugas, Surat Keterangan, Surat Keputusan, Surat Undangan, Surat Edaran, Nota Dinas, dst) dengan isi bebas, plus Buku Agenda Surat Keluar otomatis.",
                langkah=[
                    "Pilih jenis surat umum yang ingin dibuat.",
                    "Ketik isi surat secara bebas sesuai kebutuhan — kop surat, nomor otomatis, dan blok tanda tangan tetap terisi otomatis dari Pengaturan Identitas & Branding.",
                    "Simpan/cetak — surat otomatis tercatat di Buku Agenda Surat Keluar.",
                ],
                tips=[
                    "Nomor surat otomatis mengikuti pengaturan format yang sama dengan Cetak Surat Tugas Akhir, supaya penomoran tetap konsisten satu prodi."
                ],
            ),
        ],
    ),
    (
        "⚙️ Administrasi",
        "Pengaturan yang mendefinisikan identitas program studi itu sendiri — jarang berubah, biasanya hanya disentuh Kaprodi/Admin Prodi.",
        [
            dict(
                judul="Manajemen Pengguna",
                endpoint="pengaturan.pengguna",
                ringkasan="Direktori kontak pengguna/operator aplikasi (bukan sistem akun personal — SIMPRODI memakai satu password admin bersama).",
                langkah=[
                    "Tambah/ubah data kontak pengguna yang mengoperasikan SIMPRODI di prodi Anda."
                ],
                tips=[],
            ),
            dict(
                judul="Tahun Akademik",
                endpoint="pengaturan.tahun_akademik",
                ringkasan="Pengaturan tahun akademik yang sedang berjalan — tampil di topbar semua halaman.",
                langkah=[
                    "Tetapkan/ubah tahun akademik aktif di sini setiap pergantian semester/tahun."
                ],
                tips=[
                    "Banyak modul (Jadwal, Semester Pendek, dst) menyaring data berdasarkan tahun akademik aktif ini."
                ],
            ),
            dict(
                judul="Identitas & Branding",
                endpoint="pengaturan.branding",
                ringkasan="Nama institusi, fakultas, program studi, dan logo — dipakai di topbar, halaman login, dan kop surat otomatis.",
                langkah=[
                    "Isi/ubah nama institusi, fakultas, dan program studi.",
                    "Unggah logo — akan tampil di sidebar, halaman login, dan kop surat.",
                ],
                tips=[],
            ),
            dict(
                judul="Pejabat Struktural",
                endpoint="pengaturan.pejabat",
                ringkasan="Direktori Rektor, Dekan, Kaprodi, dan pejabat struktural lain — sumber blok tanda tangan SK & Surat.",
                langkah=[
                    "Tambahkan setiap pejabat (jabatan, nama lengkap dengan gelar, unit, NIP/NIDN, no. SK pengangkatan).",
                    "Klik \"Jadikan Default\" pada pejabat yang akan tampil otomatis sebagai penandatangan di SK Tugas Akhir & Surat Umum.",
                ],
                tips=[
                    "Kalau jabatan yang dicari tidak ada di daftar saran, tetap bisa diketik bebas — daftar saran hanya mempercepat pengisian.",
                    "Ganti pejabat default kapan saja (mis. pergantian Kaprodi) tanpa mengedit surat yang sudah pernah dibuat sebelumnya.",
                ],
            ),
            dict(
                judul="Import & Export Data",
                endpoint="pengaturan.import_export",
                ringkasan="Impor/ekspor data per tabel dalam format Excel.",
                langkah=[
                    "Untuk ekspor: pilih tabel/data yang ingin diunduh.",
                    "Untuk impor: unduh dulu template Excel yang sesuai, isi datanya, lalu unggah kembali ke sistem.",
                ],
                tips=[
                    "Fitur ini per-tabel, bukan cadangan menyeluruh aplikasi — untuk cadangan seluruh database, gunakan Backup & Restore di grup Pengaturan."
                ],
            ),
        ],
    ),
    (
        "⚙️ Pengaturan",
        "Hal-hal yang berhubungan dengan cara memakai aplikasi sehari-hari: panduan ini sendiri, ubah password, dan log aktivitas.",
        [
            dict(
                judul="Panduan Penggunaan",
                endpoint="panduan.index",
                ringkasan="Halaman yang sedang Anda buka sekarang — kumpulan panduan lengkap cara memakai setiap modul SIMPRODI.",
                langkah=[
                    "Klik grup pada daftar di kiri/atas untuk membuka & menutup (expand/collapse) penjelasan tiap modul di dalamnya."
                ],
                tips=[],
            ),
            dict(
                judul="Ubah Password",
                endpoint="pengaturan.password",
                ringkasan="Mengubah password login bersama SIMPRODI.",
                langkah=[
                    "Masukkan password lama, lalu password baru dua kali untuk konfirmasi.",
                    "Simpan — password baru langsung berlaku untuk login berikutnya.",
                ],
                tips=[
                    "Karena SIMPRODI memakai satu password bersama, ubah password di sini akan memengaruhi seluruh pengguna aplikasi ini."
                ],
            ),
            dict(
                judul="Log Aktivitas",
                endpoint="mutu.index",
                ringkasan="Riwayat aktivitas (tambah/ubah/hapus data) yang tercatat otomatis di hampir seluruh modul SIMPRODI.",
                langkah=[
                    "Buka tab Log untuk melihat riwayat aktivitas terbaru; gunakan filter tanggal/modul bila tersedia."
                ],
                tips=[],
            ),
            dict(
                judul="Backup & Restore",
                endpoint="backup.index",
                ringkasan="Cadangkan/pulihkan seluruh database SIMPRODI (satu file .db) langsung dari aplikasi.",
                langkah=[
                    'Klik "Backup Sekarang" untuk membuat cadangan penuh database saat ini.',
                    'Untuk memulihkan, pilih salah satu file backup dari daftar lalu klik "Restore" — database berjalan akan digantikan seluruhnya.',
                ],
                tips=[
                    "Backup lama dibersihkan otomatis (retensi 30 hari, minimal 3 file terbaru selalu disisakan) tiap aplikasi start — tidak perlu dibersihkan manual.",
                    "Kalau belum pernah backup atau sudah lebih dari 7 hari, Dashboard & Notifikasi akan menampilkan pengingat otomatis.",
                    "Setiap backup ditandai kode tahun ajaran yang sedang berjalan saat backup dibuat — gunakan dropdown di daftar Riwayat Backup untuk menyaring, misalnya mencari backup dari tahun ajaran tertentu. Ini hanya menyaring tampilan; setiap file backup tetap berisi keseluruhan data aplikasi.",
                ],
            ),
            dict(
                judul="Tentang Aplikasi",
                endpoint="tentang.index",
                ringkasan="Ringkasan versi aplikasi, identitas instansi, status database, dan status backup dalam satu halaman.",
                langkah=[
                    "Buka menunya langsung — halaman ini murni informasi, tidak ada yang perlu diisi."
                ],
                tips=[
                    "Cakupan modul (aktif vs direncanakan) di halaman ini dihitung otomatis dari modul yang sungguh terdaftar, jadi selalu sesuai keadaan aplikasi sebenarnya."
                ],
            ),
            dict(
                judul="Preferensi Tampilan",
                endpoint="preferensi.index",
                ringkasan="Kepadatan tabel, mode default grup sidebar, dan rentang hari Agenda Mendatang — berlaku di seluruh aplikasi.",
                langkah=[
                    'Pilih Kepadatan Tabel: "Nyaman" (baku) atau "Padat" untuk melihat lebih banyak baris tanpa scroll.',
                    "Pilih Mode Default Grup Sidebar: ingat pilihan terakhir, selalu buka semua, atau selalu tutup semua.",
                    "Pilih rentang hari untuk widget Agenda Mendatang di Dashboard & Pusat Notifikasi, lalu Simpan.",
                ],
                tips=[
                    "Perubahan langsung terlihat di seluruh halaman setelah disimpan, tanpa perlu restart aplikasi."
                ],
            ),
            dict(
                judul="Pusat Notifikasi",
                endpoint="notifikasi.index",
                ringkasan="Kumpulan reminder & peringatan dari seluruh modul (SDM, Tri Dharma, Kerja Sama, Mutu/AMI, Agenda Kalender, dan peringatan operasional) di satu tempat, dengan ambang hari yang bisa diatur.",
                langkah=[
                    "Gunakan filter Level (Lewat Tenggat/Segera/Info) dan Kategori untuk menyaring daftar.",
                    'Klik "Buka" pada suatu item untuk langsung menuju modul terkait.',
                    'Atur "Ambang Hari per Kategori" di bagian bawah untuk mengubah kapan suatu item mulai dianggap "Segera", lalu Simpan.',
                ],
                tips=[
                    "Ikon lonceng di topbar menampilkan badge jumlah item yang perlu perhatian (lewat tenggat + segera) di halaman manapun."
                ],
            ),
            dict(
                judul="Tema Tampilan",
                endpoint="tema.index",
                ringkasan="Pilihan aksen warna aplikasi (Indigo/Emerald/Ocean/Amber/Rose/Slate) di luar warna indigo baku SIMPRODI.",
                langkah=[
                    "Klik salah satu kartu warna — tema langsung tersimpan & diterapkan ke seluruh halaman."
                ],
                tips=[
                    "Hanya mengganti aksen warna; kontras teks & keterbacaan tabel/rekap tetap sama di semua tema (bukan mode gelap)."
                ],
            ),
        ],
    ),
]

# Catatan roadmap (Audit Lanjutan 3): Preferensi Tampilan, Pusat Notifikasi
# & Tema Tampilan sudah jadi modul nyata (lihat 3 entri terakhir grup
# "⚙️ Pengaturan" di atas) — tidak ada lagi menu roadmap/placeholder yang
# tersisa di sidebar saat ini, jadi catatan ini dikosongkan (bukan
# dihapus) supaya `roadmap_note` tetap tersedia untuk template & mudah
# diisi lagi kalau ada modul roadmap baru di kemudian hari.
PANDUAN_ROADMAP_NOTE = ""


@bp.route("/")
def index():
    return render_template(
        "panduan.html",
        panduan_groups=PANDUAN_GROUPS,
        roadmap_note=PANDUAN_ROADMAP_NOTE,
    )


# ---------------------------------------------------------------------------
# Unduh sebagai dokumen resmi (.docx) — melengkapi tombol "Unduh sebagai PDF"
# (window.print(), tetap dipertahankan sebagai jalan pintas cepat) dengan
# dokumen Word SUNGGUHAN yang tersusun sistematis: sampul, identitas
# dokumen, kata pengantar, daftar isi otomatis (field TOC bawaan Word),
# lalu seluruh isi panduan per grup/modul dengan heading berjenjang
# (Heading 1/2/3) — supaya layak dijadikan dokumen resmi yang diarsipkan,
# dicetak, atau dilampirkan ke borang akreditasi, bukan sekadar cetakan
# tampilan layar. Memakai python-docx, mengikuti pola & helper yang sama
# dengan `routes/surat.py`/`routes/surat_umum.py` (kop surat dari
# Pengaturan Identitas & Branding, generator docx via io.BytesIO).
# ---------------------------------------------------------------------------


def _toc_field(doc):
    """Sisipkan field code TOC (Table of Contents) bawaan Microsoft Word.
    Berbeda dari daftar isi manual (teks statis yang bisa basi begitu ada
    modul baru ditambahkan), field ini otomatis mengambil SELURUH heading
    berjenjang (Heading 1-3) beserta nomor halaman yang benar begitu
    pengguna membuka dokumen di Word dan menekan F9 / klik kanan ->
    "Update Field" (Word menampilkan pesan itu saat pertama dibuka —
    ini perilaku standar Word untuk field TOC, bukan bug)."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    r = run._r

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    fld_begin.set(qn("w:dirty"), "true")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    placeholder = OxmlElement("w:t")
    placeholder.text = (
        "Daftar isi belum ter-update — klik kanan di sini lalu pilih "
        '"Update Field" (atau tekan F9) untuk menampilkan nomor halaman.'
    )

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r.append(fld_begin)
    r.append(instr)
    r.append(fld_sep)
    r.append(placeholder)
    r.append(fld_end)


def _page_number_field(paragraph, kind="PAGE"):
    """Sisipkan field code PAGE atau NUMPAGES — dipakai di footer supaya
    nomor halaman "Halaman X dari Y" ikut terhitung otomatis oleh Word,
    bukan angka statis yang salah begitu isi panduan bertambah."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    run = paragraph.add_run()
    r = run._r
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = kind
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    r.append(fld_begin)
    r.append(instr)
    r.append(fld_end)


def _bangun_docx_panduan(conn):
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Cm, Pt, RGBColor

    nama_institusi = _db.get_setting(conn, "nama_institusi", "")
    nama_fakultas = _db.get_setting(conn, "nama_fakultas", "")
    nama_prodi = _db.get_setting(conn, "nama_prodi", "")
    logo_path = _db.get_setting(conn, "logo_path", "")
    tanggal_dibuat = date.today().strftime("%d %B %Y")
    total_modul = sum(len(mods) for _l, _d, mods in PANDUAN_GROUPS)

    doc = docx.Document()

    # Properti dokumen (metadata file, tampil di "Properties" Word/Explorer)
    doc.core_properties.title = f"Panduan Penggunaan {C.APP_NAME}"
    doc.core_properties.subject = "Dokumen Resmi Panduan Penggunaan Aplikasi"
    doc.core_properties.author = nama_institusi or C.APP_SHORT_NAME
    doc.core_properties.comments = (
        f"Dibuat otomatis oleh modul Panduan Penggunaan {C.APP_SHORT_NAME} v{C.APP_VERSION}."
    )

    # Margin standar dokumen resmi (2.5cm semua sisi — mengikuti pola surat
    # resmi Indonesia, sama seperti dipakai di surat.py/surat_umum.py)
    for sec in doc.sections:
        sec.top_margin = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3)
        sec.right_margin = Cm(2.5)

    # ----------------------------------------------------------- Sampul
    if (
        logo_path
        and os.path.exists(logo_path)
        and logo_path.lower().endswith((".png", ".jpg", ".jpeg"))
    ):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p_logo.add_run().add_picture(logo_path, width=Cm(2.8))
        except Exception:
            pass  # logo korup/format tak didukung -> sampul tetap tanpa logo, bukan gagal total

    for i in range(6):
        doc.add_paragraph("")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("PANDUAN PENGGUNAAN APLIKASI")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x2A, 0x2E, 0x45)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(f"{C.APP_NAME} ({C.APP_SHORT_NAME})")
    run2.bold = True
    run2.font.size = Pt(16)

    doc.add_paragraph("")
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run(f"Versi Aplikasi {C.APP_VERSION}").font.size = Pt(11)

    for i in range(8):
        doc.add_paragraph("")

    if nama_institusi:
        p4 = doc.add_paragraph()
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r4 = p4.add_run(nama_institusi)
        r4.bold = True
        r4.font.size = Pt(13)
    if nama_fakultas or nama_prodi:
        p5 = doc.add_paragraph()
        p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p5.add_run(" | ".join(x for x in [nama_fakultas, nama_prodi] if x)).font.size = Pt(11)
    p6 = doc.add_paragraph()
    p6.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p6.add_run(f"Dicetak pada {tanggal_dibuat}").font.size = Pt(10)

    doc.add_page_break()

    # --------------------------------------------------- Identitas Dokumen
    doc.add_heading("Identitas Dokumen", level=1)
    identitas = [
        ("Nama Dokumen", "Panduan Penggunaan Aplikasi"),
        ("Nama Aplikasi", f"{C.APP_NAME} ({C.APP_SHORT_NAME})"),
        ("Versi Aplikasi", C.APP_VERSION),
        ("Instansi", nama_institusi or "-"),
        ("Program Studi", nama_prodi or "-"),
        ("Jumlah Grup Modul", str(len(PANDUAN_GROUPS))),
        ("Jumlah Modul Dijelaskan", str(total_modul)),
        ("Tanggal Dokumen Dibuat", tanggal_dibuat),
        (
            "Sumber",
            "Dibangkitkan otomatis dari halaman Panduan Penggunaan "
            "di dalam aplikasi — selalu sinkron dengan menu sidebar terbaru.",
        ),
    ]
    table = doc.add_table(rows=0, cols=2)
    table.style = (
        "Light Grid Accent 1"
        if "Light Grid Accent 1" in [s.name for s in doc.styles]
        else "Table Grid"
    )
    for k, v in identitas:
        cells = table.add_row().cells
        cells[0].text = k
        cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = v

    doc.add_paragraph("")

    # ---------------------------------------------------- Kata Pengantar
    doc.add_heading("Kata Pengantar", level=1)
    doc.add_paragraph(
        f"Dokumen ini disusun sebagai panduan resmi penggunaan aplikasi "
        f"{C.APP_NAME} ({C.APP_SHORT_NAME}) bagi Ketua Program Studi/operator "
        f"yang mengelola data akademik, kemahasiswaan, sumber daya manusia, "
        f"tri dharma, kerja sama, mutu, dan administrasi program studi. "
        f"Setiap modul dijelaskan secara ringkas: fungsi modul, langkah "
        f"pemakaian berurutan, dan catatan/tips praktis bila relevan."
    )
    doc.add_paragraph(
        "Susunan dokumen ini mengikuti pengelompokan menu pada sidebar "
        "aplikasi, sehingga urutan panduan langsung sejalan dengan urutan "
        "menu yang dilihat pengguna sehari-hari. Dokumen ini dibangkitkan "
        "otomatis dari konten yang sama dengan yang tampil di halaman "
        "Panduan Penggunaan dalam aplikasi, sehingga isinya senantiasa "
        "sinkron dengan modul yang benar-benar tersedia."
    )

    doc.add_page_break()

    # -------------------------------------------------------- Daftar Isi
    doc.add_heading("Daftar Isi", level=1)
    doc.add_paragraph(
        "Daftar isi berikut memakai fitur Table of Contents bawaan "
        "Microsoft Word. Jika nomor halaman belum tampil, klik kanan pada "
        'area di bawah ini lalu pilih "Update Field", atau tekan F9.'
    ).italic = True
    _toc_field(doc)

    doc.add_page_break()

    # -------------------------------------------------------- Isi Panduan
    doc.add_heading("Isi Panduan per Modul", level=1)
    for gi, (group_label, group_desc, modules) in enumerate(PANDUAN_GROUPS, start=1):
        doc.add_heading(f"{gi}. {group_label}", level=1)
        doc.add_paragraph(group_desc)
        for mi, m in enumerate(modules, start=1):
            doc.add_heading(f"{gi}.{mi} {m['judul']}", level=2)
            doc.add_paragraph(m["ringkasan"])
            if m.get("langkah"):
                doc.add_heading("Langkah Pemakaian", level=3)
                for langkah in m["langkah"]:
                    doc.add_paragraph(langkah, style="List Number")
            if m.get("tips"):
                doc.add_heading("Tips", level=3)
                for tip in m["tips"]:
                    doc.add_paragraph(tip, style="List Bullet")

    # ------------------------------------------------------------- Penutup
    doc.add_heading("Catatan Penutup", level=1)
    doc.add_paragraph(PANDUAN_ROADMAP_NOTE)
    doc.add_paragraph(
        "Dokumen ini dapat diunduh ulang kapan saja dari menu Panduan "
        "Penggunaan (grup Pengaturan) dan akan selalu mencerminkan daftar "
        "modul aplikasi versi terkini pada saat diunduh."
    )

    # ------------------------------------------------- Header & Footer
    section = doc.sections[0]
    header_p = section.header.paragraphs[0]
    header_p.text = f"Panduan Penggunaan — {C.APP_NAME} ({C.APP_SHORT_NAME})"
    header_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_p.runs[0].font.size = Pt(9)
    header_p.runs[0].font.color.rgb = RGBColor(0x77, 0x77, 0x77)

    footer_p = section.footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.add_run("Halaman ").font.size = Pt(9)
    _page_number_field(footer_p, "PAGE")
    footer_p.add_run(" dari ").font.size = Pt(9)
    _page_number_field(footer_p, "NUMPAGES")

    return doc


@bp.route("/unduh")
def unduh():
    """Unduh dokumen resmi (.docx) — lihat catatan panjang di
    `_bangun_docx_panduan()`. Endpoint ini murni baca (tidak menulis apa
    pun ke database), aman dipanggil berulang kali."""
    conn = current_app.get_db()
    doc = _bangun_docx_panduan(conn)
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    tanggal = date.today().strftime("%Y%m%d")
    return send_file(
        buf,
        as_attachment=True,
        download_name=f"Panduan-Penggunaan-{C.APP_SHORT_NAME}-{tanggal}.docx",
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
