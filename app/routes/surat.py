# -*- coding: utf-8 -*-
"""routes/surat.py — Cetak Surat Tugas Akhir (SK Pembimbing, SK Yudisium
per-mahasiswa & per-Tahap/Gelombang, Undangan Seminar/Sidang) sebagai
dokumen Word (.docx) siap-cetak, terisi otomatis dari data yang
tersimpan. Logika isi surat porting dari CetakSuratView versi desktop
(views_premium.py) — hanya cara menyimpannya yang beda: langsung diunduh
via browser, bukan dialog "Save As" native.

Blok kop (`_header`) & tanda tangan (`_footer_ttd`) senada gaya dengan
`routes/surat_umum.py` (Modul 8 — Generator Surat Umum di luar Tugas
Akhir); keduanya sama-sama membaca penandatangan default dari
Pengaturan > Pejabat Struktural.

Audit Lanjutan 10 — kerapian modul ini: `import docx` & submodulnya
SEBELUMNYA diimpor ulang secara identik di 5 fungsi berbeda (padahal
blueprint ini SELALU dimuat penuh saat app start lewat
`app/__init__.py::create_app()`, jadi tidak ada penghematan startup yang
didapat dari menunda impor per-fungsi seperti di modul lain yang docx-
nya cuma dipakai di satu titik) — sekarang diimpor sekali di sini."""

import io

import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from flask import (
    Blueprint,
    current_app,
    flash,
    redirect,
    render_template,
    request,
    send_file,
    url_for,
)

from app import db as _db
from app import error_utils as EH
from app import logic as L

bp = Blueprint("surat", __name__, url_prefix="/surat")


@bp.route("/")
def index():
    conn = current_app.get_db()
    mhs_list = conn.execute("SELECT id, nim, nama FROM mahasiswa ORDER BY nama").fetchall()
    return render_template("surat.html", mhs_list=mhs_list, jenis_list=JENIS_SURAT)


def _header(doc, conn):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(_db.get_setting(conn, "nama_institusi"))
    run.bold = True
    run.font.size = Pt(13)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run(
        f"{_db.get_setting(conn, 'nama_fakultas')} | " f"{_db.get_setting(conn, 'nama_prodi')}"
    ).font.size = Pt(10)
    doc.add_paragraph("=" * 70).alignment = WD_ALIGN_PARAGRAPH.CENTER


def _footer_ttd(doc, conn):
    """Blok tanda tangan SK Tugas Akhir/Undangan. SEBELUM perbaikan ini,
    jabatan & nama penandatangan SELALU hardcode "Ketua Program Studi,"
    dengan baris nama kosong — 2 pengaturan yang sudah ada
    (nama_penandatangan_default/jabatan_penandatangan_default, dipakai
    normal di Surat Umum) sama sekali tidak dibaca di sini, jadi Kaprodi
    harus menulis tangan nama penandatangan di SETIAP SK yang dicetak.
    Sekarang dibaca dari Pengaturan > Pejabat Struktural (via
    "Jadikan Default Penandatangan", yang menulis ke 3 kunci pengaturan
    ini) — kalau belum pernah diisi, jatuh balik ke perilaku lama supaya
    tidak ada baris kosong yang membingungkan di dokumen."""
    jabatan = _db.get_setting(conn, "jabatan_penandatangan_default", "Ketua Program Studi")
    nama = _db.get_setting(conn, "nama_penandatangan_default", "")
    nip_nidn = _db.get_setting(conn, "nip_nidn_penandatangan_default", "")
    doc.add_paragraph("")
    doc.add_paragraph(f"{jabatan or 'Ketua Program Studi'},")
    doc.add_paragraph("")
    doc.add_paragraph("")
    doc.add_paragraph(nama or "(...........................................)")
    if nip_nidn:
        doc.add_paragraph(f"NIP/NIDN. {nip_nidn}")


def _gaya_tabel(doc):
    """Nama gaya tabel bergaris yang dipakai konsisten di semua tabel yang
    dibuat modul ini (SK Pembimbing/Yudisium, tabel rekap per-Tahap,
    Undangan). SEBELUM perbaikan ini, `doc.add_table()` dipakai di 4
    tempat tanpa `table.style` sama sekali — jatuh ke gaya bawaan
    python-docx "Normal Table", yang TIDAK PUNYA garis batas sel apa pun
    (diverifikasi: XML gaya itu tidak berisi definisi `tblBorders`).
    Hasilnya, SK/Undangan yang dicetak terlihat seperti dua kolom teks
    rata kiri tanpa kotak pembatas — bukan tabel data formal yang lazim
    di surat resmi institusi. Dipakai gaya yang sama dengan tabel di
    `routes/panduan.py` (dokumen Word lain yang sudah rapi) supaya
    konsisten: "Light Grid Accent 1" kalau tersedia di tema dokumen,
    fallback ke "Table Grid" bawaan (selalu tersedia)."""
    return "Light Grid Accent 1" if "Light Grid Accent 1" in [s.name for s in doc.styles] else "Table Grid"


def _tabel(doc, pairs):
    table = doc.add_table(rows=0, cols=2)
    table.style = _gaya_tabel(doc)
    for k, v in pairs:
        cells = table.add_row().cells
        cells[0].text = k
        if cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True
        cells[1].text = str(v) if v is not None else "-"
    return table


def _kirim_docx(doc, filename):
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return send_file(
        buf,
        as_attachment=True,
        download_name=filename,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


def _gen_sk_pembimbing(conn, mid):
    m = conn.execute("SELECT * FROM mahasiswa WHERE id=?", (mid,)).fetchone()
    pp = conn.execute(
        "SELECT pp.*, d1.nama AS p1, d2.nama AS p2 FROM penetapan_pembimbing pp "
        "LEFT JOIN dosen d1 ON d1.id=pp.pembimbing1_id "
        "LEFT JOIN dosen d2 ON d2.id=pp.pembimbing2_id WHERE pp.mahasiswa_id=?",
        (mid,),
    ).fetchone()
    if not pp:
        return None
    doc = docx.Document()
    _header(doc, conn)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SURAT KEPUTUSAN\nPENETAPAN DOSEN PEMBIMBING SKRIPSI")
    r.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph(f"Nomor: {pp['no_sk'] or '.......................'}")
    doc.add_paragraph("")
    doc.add_paragraph("Menetapkan mahasiswa berikut untuk dibimbing dalam penyusunan Skripsi:")
    _tabel(
        doc,
        [
            ("NIM", m["nim"]),
            ("Nama", m["nama"]),
            ("Semester", pp["semester"] or "-"),
            ("Tahap/Gelombang", pp["tahap"] or "-"),
            ("Judul Skripsi", pp["judul_final"] or "-"),
            ("Pembimbing 1", pp["p1"] or "-"),
            ("Pembimbing 2", pp["p2"] or "-"),
            ("Tanggal Penetapan", pp["tgl_penetapan"] or "-"),
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph(f"Tanggal: {pp['tgl_penetapan'] or '.......................'}")
    _footer_ttd(doc, conn)
    return doc, f"SK_Pembimbing_{m['nim']}.docx"


def _gen_sk_yudisium(conn, mid):
    m = conn.execute("SELECT * FROM mahasiswa WHERE id=?", (mid,)).fetchone()
    rows = L.rencana_yudisium_rows(conn)
    row = next((r for r in rows if r["mahasiswa_id"] == mid), None)
    if not row:
        return None
    doc = docx.Document()
    _header(doc, conn)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("SURAT KEPUTUSAN\nYUDISIUM KELULUSAN")
    r.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph(f"Nomor: {row['no_sk'] or '.......................'}")
    doc.add_paragraph("")
    _tabel(
        doc,
        [
            ("NIM", row["nim"]),
            ("Nama", row["nama"]),
            ("Judul Skripsi", row["judul_sidang"] or "-"),
            ("Nilai Angka", row["nilai_angka"]),
            ("Nilai Huruf", row["nilai_huruf"]),
            ("IPK Final", row["ipk_final"]),
            ("Predikat Kelulusan", row["predikat"]),
            ("Tanggal Yudisium", row["tgl_yudisium"] or "-"),
        ],
    )
    _footer_ttd(doc, conn)
    return doc, f"SK_Yudisium_{m['nim']}.docx"


def _gen_sk_yudisium_tahap(conn, tahap):
    """SK Yudisium untuk SATU TAHAP/GELOMBANG sekaligus — lazimnya satu SK
    Yudisium memang berlaku untuk satu batch kelulusan (bukan satu SK per
    mahasiswa satu-satu seperti _gen_sk_yudisium di atas, yang tetap
    dipertahankan untuk kebutuhan cetak-ulang per-individu). No. SK &
    Tgl Yudisium diambil dari data yang sudah ditetapkan lewat
    'kelulusan.yudisium_tetapkan_tahap' (kolom no_sk/tgl_yudisium di baris
    yudisium — di-set sama untuk semua mahasiswa dalam tahap tsb)."""
    rows = L.rencana_yudisium_rows(conn, tahap)
    if not rows:
        return None
    no_sk = next((r["no_sk"] for r in rows if r["no_sk"]), "") or "......................."
    tgl = next((r["tgl_yudisium"] for r in rows if r["tgl_yudisium"]), "") or "......................."

    doc = docx.Document()
    _header(doc, conn)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = title.add_run(f"SURAT KEPUTUSAN\nYUDISIUM KELULUSAN\nTAHAP/GELOMBANG: {tahap.upper()}")
    r0.bold = True
    r0.font.size = Pt(12)
    doc.add_paragraph(f"Nomor: {no_sk}")
    doc.add_paragraph(f"Tanggal Yudisium: {tgl}")
    doc.add_paragraph("")
    doc.add_paragraph(
        f"Menetapkan {len(rows)} mahasiswa berikut LULUS yudisium pada Tahap/Gelombang "
        f'"{tahap}":'
    )
    table = doc.add_table(rows=1, cols=6)
    table.style = _gaya_tabel(doc)
    hdr = table.rows[0].cells
    for i, judul in enumerate(["No", "NIM", "Nama", "Nilai Huruf", "IPK Final", "Predikat"]):
        hdr[i].text = judul
    for i, row in enumerate(rows, start=1):
        cells = table.add_row().cells
        cells[0].text = str(i)
        cells[1].text = row["nim"]
        cells[2].text = row["nama"]
        cells[3].text = str(row["nilai_huruf"] or "-")
        cells[4].text = str(row["ipk_final"]) if row["ipk_final"] is not None else "-"
        cells[5].text = row["predikat"] or "-"
    _footer_ttd(doc, conn)
    nama_file = tahap.replace(" ", "_").replace("/", "-")
    return doc, f"SK_Yudisium_Tahap_{nama_file}.docx"


def _gen_undangan(conn, mid, jenis):
    m = conn.execute("SELECT * FROM mahasiswa WHERE id=?", (mid,)).fetchone()
    if jenis == "Seminar":
        row = conn.execute(
            "SELECT s.*, pp.judul_final FROM seminar s "
            "LEFT JOIN penetapan_pembimbing pp ON pp.mahasiswa_id=s.mahasiswa_id "
            "WHERE s.mahasiswa_id=? ORDER BY s.id DESC LIMIT 1",
            (mid,),
        ).fetchone()
        tim = (
            []
            if not row
            else [
                ("Ketua Penguji", L.dosen_nama(conn, row["penguji_ketua_id"])),
                ("Anggota Penguji 1", L.dosen_nama(conn, row["penguji_anggota1_id"])),
                ("Anggota Penguji 2", L.dosen_nama(conn, row["penguji_anggota2_id"])),
            ]
        )
        tgl, jam = (row["tgl_seminar"], row["jam"]) if row else ("-", "-")
        ruang = L.ruangan_nama(conn, row["ruangan_id"]) if row else ""
        # Audit Lanjutan 10 (kerapian): disamakan gaya penulisannya dgn
        # cabang Sidang di bawah -- ekspresi "(judul_final if row else '')
        # or ''" sebelumnya tidak konsisten dgn cabang Sidang yg memakai
        # bentuk "(A or B) if row else ''"; perilaku sama, cuma dirapikan
        # supaya kedua cabang mudah dibandingkan sekilas.
        judul = (row["judul_final"] or "") if row else ""
    else:
        row = conn.execute(
            "SELECT sd.*, pp.judul_final FROM sidang sd "
            "LEFT JOIN penetapan_pembimbing pp ON pp.mahasiswa_id=sd.mahasiswa_id "
            "WHERE sd.mahasiswa_id=? ORDER BY sd.id DESC LIMIT 1",
            (mid,),
        ).fetchone()
        tim = (
            []
            if not row
            else [
                ("Ketua", L.dosen_nama(conn, row["ketua_id"])),
                ("Sekretaris", L.dosen_nama(conn, row["sekretaris_id"])),
                ("Anggota 1", L.dosen_nama(conn, row["anggota1_id"])),
                ("Anggota 2", L.dosen_nama(conn, row["anggota2_id"])),
                ("Anggota 3", L.dosen_nama(conn, row["anggota3_id"])),
            ]
        )
        tgl, jam = (row["tgl_sidang"], row["jam_sidang"]) if row else ("-", "-")
        ruang = L.ruangan_nama(conn, row["ruangan_id"]) if row else ""
        judul = ((row["judul_sidang"] or row["judul_final"]) or "") if row else ""

    if not row:
        return None

    doc = docx.Document()
    _header(doc, conn)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(f"UNDANGAN {jenis.upper()} SKRIPSI")
    r.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph("")
    doc.add_paragraph(
        f"Dengan hormat, mengundang Bapak/Ibu untuk hadir sebagai tim penguji pada "
        f"pelaksanaan {jenis} Skripsi mahasiswa berikut:"
    )
    _tabel(
        doc,
        [
            ("NIM", m["nim"]),
            ("Nama", m["nama"]),
            ("Judul", judul or "-"),
            ("Hari/Tanggal", tgl or "-"),
            ("Jam", jam or "-"),
            ("Ruangan", ruang or "-"),
        ],
    )
    doc.add_paragraph("")
    doc.add_paragraph("Susunan Tim:")
    t2 = doc.add_table(rows=0, cols=2)
    t2.style = _gaya_tabel(doc)
    for peran, nama in tim:
        cells = t2.add_row().cells
        cells[0].text = peran
        cells[1].text = nama or "-"
    doc.add_paragraph("")
    doc.add_paragraph("Atas kehadiran dan kerja sama Bapak/Ibu, kami ucapkan terima kasih.")
    _footer_ttd(doc, conn)
    return doc, f"Undangan_{jenis}_{m['nim']}.docx"


_GENERATORS = {
    "SK Pembimbing": _gen_sk_pembimbing,
    "SK Yudisium": _gen_sk_yudisium,
    "Undangan Seminar": lambda conn, mid: _gen_undangan(conn, mid, "Seminar"),
    "Undangan Sidang": lambda conn, mid: _gen_undangan(conn, mid, "Sidang"),
}

_PESAN_KOSONG = {
    "SK Pembimbing": "Mahasiswa ini belum memiliki data Penetapan Pembimbing.",
    "SK Yudisium": "Mahasiswa ini belum memiliki data Rencana Yudisium (harus LULUS Sidang dahulu).",
    "Undangan Seminar": "Mahasiswa ini belum memiliki jadwal Seminar.",
    "Undangan Sidang": "Mahasiswa ini belum memiliki jadwal Sidang.",
}

# Audit Lanjutan 10 (kerapian): SEBELUM perbaikan ini, daftar 4 jenis
# surat tertulis TIGA KALI secara terpisah dan manual — sekali sebagai
# `JENIS_SURAT` (dipakai buat dropdown di index()), dan sekali lagi
# sebagai key di masing-masing `_GENERATORS` & `_PESAN_KOSONG` di atas.
# Risikonya: kalau suatu saat ada jenis surat baru ditambahkan, mudah
# lupa memperbarui salah satu dari ketiganya (mis. dropdown bertambah
# tapi _PESAN_KOSONG lupa diisi -> pesan error generik yang kurang
# informatif; atau sebaliknya). Sekarang `JENIS_SURAT` diturunkan
# otomatis dari `_GENERATORS`, satu sumber kebenaran.
JENIS_SURAT = list(_GENERATORS)


@bp.route("/sk-yudisium-tahap")
def sk_yudisium_tahap():
    """Cetak SATU dokumen SK Yudisium untuk seluruh mahasiswa pada satu
    Tahap/Gelombang (lihat _gen_sk_yudisium_tahap). Dipanggil dari tombol
    'Cetak SK Yudisium (Tahap Ini)' di halaman Rencana Yudisium."""
    conn = current_app.get_db()
    tahap = request.args.get("tahap", "").strip()
    if not tahap or tahap == "Semua":
        flash("Pilih Tahap/Gelombang tertentu dahulu untuk mencetak SK Yudisium per tahap.", "error")
        return redirect(url_for("kelulusan.yudisium_list"))
    # Audit Lanjutan 10 — SEBELUM perbaikan ini, seluruh alur di bawah
    # (pembuatan dokumen .docx + `_db.log`) TIDAK dibungkus try/except,
    # beda dgn pola seragam di hampir semua handler lain di aplikasi ini
    # (lihat `app/error_utils.py`). Galat tak terduga (mis. `_db.log`
    # gagal commit krn basis data sedang dipakai proses lain) akan bocor
    # jadi halaman 500 Flask mentah alih-alih flash pesan ramah.
    try:
        hasil = _gen_sk_yudisium_tahap(conn, tahap)
        if not hasil:
            flash(
                f'Belum ada mahasiswa LULUS pada tahap "{tahap}" untuk dibuatkan SK Yudisium.',
                "error",
            )
            return redirect(url_for("kelulusan.yudisium_list", tahap=tahap))
        doc, filename = hasil
        _db.log(conn, "Cetak Surat", f"SK Yudisium per Tahap - {filename}")
        return _kirim_docx(doc, filename)
    except Exception as e:
        EH.flash_gagal_simpan(e, "Gagal membuat dokumen SK Yudisium per Tahap")
        return redirect(url_for("kelulusan.yudisium_list", tahap=tahap))


@bp.route("/buat", methods=["POST"])
def buat():
    conn = current_app.get_db()
    mid = request.form.get("mahasiswa_id", type=int)
    jenis = request.form.get("jenis", "")
    if not mid or jenis not in _GENERATORS:
        flash("Pilih mahasiswa dan jenis surat terlebih dahulu.", "error")
        return redirect(url_for("surat.index"))
    # Audit Lanjutan 10 — sama seperti sk_yudisium_tahap() di atas: alur
    # generate .docx + log aktivitas sebelumnya tidak dibungkus try/except
    # sama sekali (satu-satunya route pembuat dokumen di aplikasi ini yang
    # begitu — bandingkan `simpan_kriteria`/`simpan_target_iku` dkk di
    # routes/mutu.py yang seragam pakai `EH.flash_gagal_simpan`).
    try:
        hasil = _GENERATORS[jenis](conn, mid)
        if not hasil:
            flash(_PESAN_KOSONG.get(jenis, "Data belum lengkap untuk membuat surat ini."), "error")
            return redirect(url_for("surat.index"))
        doc, filename = hasil
        _db.log(conn, "Cetak Surat", f"{jenis} - {filename}")
        return _kirim_docx(doc, filename)
    except Exception as e:
        EH.flash_gagal_simpan(e, f"Gagal membuat dokumen {jenis}")
        return redirect(url_for("surat.index"))
