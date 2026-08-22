import io
import os
import sys
import tempfile
import zipfile

sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

tmpdir = tempfile.mkdtemp()
os.environ["HOME"] = tmpdir

from app import create_app  # noqa: E402
from app import db as _db  # noqa: E402
from app import constants as C  # noqa: E402

db_path = os.path.join(tmpdir, "test.db")
app = create_app(db_path=db_path)
app.config["TESTING"] = True
app.config["WTF_CSRF_ENABLED"] = False
client = app.test_client()

FAILS = []
def check(label, cond):
    print(f"[{'OK' if cond else 'FAIL'}] {label}")
    if not cond:
        FAILS.append(label)

client.get("/login")
client.post("/login", data={"username": "kaprodi", "password1": "test1234", "password2": "test1234"}, follow_redirects=True)

# ------------------------------------------------------------- Halaman index
idx = client.get("/panduan/")
check("Halaman Panduan Penggunaan -> 200", idx.status_code == 200)
check("Tombol 'Unduh Dokumen Resmi (.docx)' ada di halaman", b"panduan.unduh" in idx.data or b"/panduan/unduh" in idx.data)
check("Tombol cetak/PDF (window.print) tetap ada", b"panduanPrint" in idx.data)

# ------------------------------------------------------------------ Unduhan
resp = client.get("/panduan/unduh")
check("GET /panduan/unduh -> 200", resp.status_code == 200)
check(
    "Content-Type file .docx sesuai OOXML Word",
    resp.headers.get("Content-Type") == "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
)
check("Header Content-Disposition menandai attachment (bukan tampil di browser)", "attachment" in resp.headers.get("Content-Disposition", ""))
check(
    f"Nama file unduhan memuat kode aplikasi ({C.APP_SHORT_NAME})",
    C.APP_SHORT_NAME in resp.headers.get("Content-Disposition", ""),
)
check("File yang dihasilkan tidak kosong (>10KB, dokumen sungguhan bukan file kosong)", len(resp.data) > 10_000)

# --------------------------------------------------- Validitas struktur file
buf = io.BytesIO(resp.data)
try:
    z = zipfile.ZipFile(buf)
    bad_file = z.testzip()
    check("File .docx adalah ZIP OOXML yang valid (tidak korup)", bad_file is None)
    check("word/document.xml ada di dalam paket", "word/document.xml" in z.namelist())
except zipfile.BadZipFile:
    check("File .docx adalah ZIP OOXML yang valid (tidak korup)", False)

# ------------------------------------------------------- Isi dokumen (python-docx)
buf.seek(0)
import docx  # noqa: E402
d = docx.Document(buf)

check("Dokumen punya isi (>100 paragraf, bukan dokumen kosong)", len(d.paragraphs) > 100)

heading_texts = [p.text for p in d.paragraphs if p.style.name.startswith("Heading")]
check("Bagian 'Identitas Dokumen' ada", "Identitas Dokumen" in heading_texts)
check("Bagian 'Kata Pengantar' ada", "Kata Pengantar" in heading_texts)
check("Bagian 'Daftar Isi' ada", "Daftar Isi" in heading_texts)
check("Bagian 'Isi Panduan per Modul' ada", "Isi Panduan per Modul" in heading_texts)
check("Bagian 'Catatan Penutup' ada", "Catatan Penutup" in heading_texts)

from app.routes.panduan import PANDUAN_GROUPS  # noqa: E402
total_modul = sum(len(mods) for _l, _d, mods in PANDUAN_GROUPS)
group_headings = [h for h in heading_texts if any(h.startswith(f"{i}. ") for i in range(1, len(PANDUAN_GROUPS) + 1))]
check(f"Seluruh {len(PANDUAN_GROUPS)} grup panduan punya heading level 1 bernomor", len(group_headings) == len(PANDUAN_GROUPS))

# Heading modul berpola "<grup>.<modul> <judul>", mis. "3.2 Seminar & Sidang"
import re  # noqa: E402
pola_modul = re.compile(r"^\d+\.\d+ ")
module_headings = [h for h in heading_texts if pola_modul.match(h)]
check(f"Seluruh {total_modul} modul punya heading level 2 bernomor (grup.modul)", len(module_headings) == total_modul)

check("Nama modul 'Kurikulum & OBE' tercantum sebagai heading", any("Kurikulum & OBE" in h for h in module_headings))
check("Nama modul 'Backup & Restore' tercantum sebagai heading", any("Backup & Restore" in h for h in module_headings))

check("Metadata judul dokumen (core_properties.title) terisi", bool(d.core_properties.title) and C.APP_NAME in d.core_properties.title)

# Field TOC (Table of Contents) harus tersemat sebagai field code Word asli,
# bukan teks statis -> cek keberadaan instrText "TOC" di XML mentah paragraf.
toc_xml = None
for p in d.paragraphs:
    if "TOC" in p._p.xml and "instrText" in p._p.xml:
        toc_xml = p._p.xml
        break
check("Field code TOC (Table of Contents) tersemat di dokumen", toc_xml is not None)

# Field PAGE/NUMPAGES di footer
footer_p = d.sections[0].footer.paragraphs[0]
footer_xml = footer_p._p.xml
check("Field nomor halaman (PAGE) tersemat di footer", "PAGE" in footer_xml and "instrText" in footer_xml)
check("Field total halaman (NUMPAGES) tersemat di footer", "NUMPAGES" in footer_xml)

# Panggil ulang -> harus tetap berhasil (murni baca, idempoten)
resp2 = client.get("/panduan/unduh")
check("Bisa diunduh berulang kali tanpa efek samping (idempoten)", resp2.status_code == 200 and len(resp2.data) > 10_000)

print("\n=== SELESAI ===")
if FAILS:
    print("ADA YANG GAGAL:", FAILS)
    sys.exit(1)
else:
    print("SEMUA TES LULUS.")
